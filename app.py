#!/usr/bin/env python3
"""buergcontrolBASE v1.0 - Die Referenz - für automatisiertes Bürgschaftscontolling"""

import streamlit as st
import pandas as pd
from datetime import timedelta, date, datetime
import io
import numpy as np
import re
from typing import List, Dict, Optional, Tuple
import json
import os
import traceback
import time
from streamlit_option_menu import option_menu
from collections import defaultdict

# Konfiguration
st.set_page_config(
    page_title="buergcontrolBASE - Bürgschaftsberechnung",
    page_icon="🏛️",
    layout="wide"
)

# buergcontolBASE Branding CSS - Komprimiert
st.markdown("""<style>[data-testid="stTooltipIcon"],[data-testid="tooltipHoverTarget"],div[role="tooltip"]{position:relative!important;z-index:99999!important}:root{--vb-primary:#8B1C1C;--vb-dark:#5C1111;--vb-light:#fee2e2;--vb-gray:#f9fafb}.stApp button:not([title*="help"]):not([role="tab"]):not(:disabled){background-color:var(--vb-primary)!important;color:white!important;border:none!important;border-radius:6px!important;font-weight:600!important}.stApp button:not([title*="help"]):not([role="tab"]):hover:not(:disabled){background-color:var(--vb-dark)!important}button[title="Show help"],button[kind="help"]{background-color:transparent!important}.stApp{background-color:#fafafa}section[data-testid="stSidebar"]{background-color:var(--vb-gray);border-right:3px solid var(--vb-primary)!important}.stSuccess{background-color:var(--vb-light)!important;color:var(--vb-primary)!important;border-left:4px solid var(--vb-primary)!important}.stInfo{background-color:var(--vb-gray)!important;border-left:4px solid var(--vb-primary)!important}[data-testid="metric-container"]{background-color:white;border:2px solid var(--vb-primary);border-radius:8px;padding:1rem}.stTabs [data-baseweb="tab-list"] button[aria-selected="true"]{border-bottom-color:var(--vb-primary)!important}a{color:var(--vb-primary)}input[type="checkbox"]:checked,input[type="radio"]:checked{accent-color:var(--vb-primary)}section[data-testid="stSidebar"] hr{border:none!important;border-top:2px solid var(--vb-primary)!important;margin:1rem 0!important}</style>""", unsafe_allow_html=True)
# Logo als HTML-Komponente
def render_logo(size="normal", with_tagline=True):
    """Rendert das buergcontolBASE Logo mit optionalem Tagline"""
    if size == "small":
        font_size = "1rem"
        padding = "6px 10px"
        tagline_size = "0.8rem"
    else:
        font_size = "1.5rem"
        padding = "10px 16px"
        tagline_size = "1rem"
    
    logo_html = f"""
    <div style="text-align: center; margin: 20px 0;">
        <div style="display: flex; align-items: center; gap: 0; justify-content: center;">
            <div style="background: #8B1C1C; padding: {padding}; border-radius: 8px; display: flex; align-items: baseline;">
                <span style="color: white; font-weight: 900; font-size: {font_size};">V</span>
                <span style="color: white; font-weight: 500; font-size: {font_size};">erwahr</span>
            </div>
            <span style="color: #8B1C1C; font-weight: 700; font-size: {font_size}; margin-left: 4px;">BASE</span>
        </div>"""
    
    if with_tagline:
        logo_html += f"""
        <div style="margin-top: 10px; text-align: center;">
            <div style="color: #8B1C1C; font-size: {tagline_size}; line-height: 1.4;">
                <div>Die Referenz - für automatisiertes</div>
                <div>Bürgschaftscontrolling</div>
            </div>
        </div>"""
    
    logo_html += "</div>"
    return logo_html

# === LOGIN UND AUTHENTIFIZIERUNG ===

def load_config():
    """Lädt die Konfiguration aus externer Datei"""
    config_file = "config.json"
    
    if os.path.exists(config_file):
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            st.error(f"❌ Fehler beim Laden der Konfiguration: {e}")
            return {
                "TEST-2024-DEMO-1234": {
                    "kunde": "Demo Kunde",
                    "logo": None,
                    "user": "test",
                    "pass": "test123"
                }
            }
    else:
        st.warning("⚠️ Keine config.json gefunden. Nutze Demo-Konfiguration.")
        return {
            "TEST-2024-DEMO-1234": {"kunde": "Test Mandant", "logo": None, "user": "test", "pass": "test123"}
        }

def validate_activation_code(code):
    """Validiert den Aktivierungscode"""
    config = load_config()
    return config.get(code)

def check_credentials(username, password, mandant_data):
    """Prüft Login-Credentials"""
    return (username == mandant_data.get('user') and 
            password == mandant_data.get('pass'))

def show_login():
    """Zeigt Login-Screen mit Aktivierungscode und Credentials"""
    if 'authenticated' not in st.session_state:
        st.session_state['authenticated'] = False
        st.session_state['mandant'] = None
        st.session_state['mandant_logo'] = None
        st.session_state['mandant_data'] = None
    
    if not st.session_state['authenticated']:
        st.markdown("""
        <style>
        .main > div {
            padding-top: 2rem;
        }
        div[data-testid="column"]:has(h1:contains("Anmeldung")) {
            background: white;
            padding: 2rem;
            border-radius: 12px;
            border: 3px solid var(--vb-primary);
            box-shadow: 0 4px 12px rgba(0,0,0,0.1);
        }
        </style>
        """, unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns([1,2,1])
        with col2:
            st.markdown("""
            <div style="background: white; padding: 2rem; border-radius: 12px; border: 3px solid #8B1C1C;">
            """ + render_logo() + """
            </div>
            """, unsafe_allow_html=True)
            st.title("Anmeldung")
            st.markdown("---")
            
            activation_code = st.text_input("Aktivierungscode", 
                                          placeholder="XXX-XXXX-XXXX-XXXX",
                                          help="Geben Sie Ihren mandantenspezifischen Aktivierungscode ein")
            username = st.text_input("Benutzername", help="Ihr Benutzername für diesen Mandanten")
            password = st.text_input("Passwort", type="password", help="Ihr Passwort")
            
            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                if st.button("🔑 Anmelden", type="primary", use_container_width=True):
                    mandant_data = validate_activation_code(activation_code)
                    if mandant_data and check_credentials(username, password, mandant_data):
                        st.session_state['authenticated'] = True
                        st.session_state['mandant'] = mandant_data['kunde']
                        st.session_state['mandant_logo'] = mandant_data.get('logo')
                        st.session_state['mandant_data'] = mandant_data
                        st.success("✅ Anmeldung erfolgreich!")
                        st.rerun()
                    else:
                        st.error("❌ Ungültige Anmeldedaten!")
            
            with col_btn2:
                if st.button("ℹ️ Demo-Zugänge", use_container_width=True):
                    st.info("""
                    **Demo-Zugänge:**
                    - Code: `TEST-2024-DEMO-1234`
                    - User: `test` / Pass: `test123`
                    """)
        
        st.stop()

# === SETTINGS MANAGEMENT ===

def load_settings():
    """Lädt mandantenspezifische Settings"""
    if 'mandant' not in st.session_state:
        return {}
    
    settings_file = f"settings_{st.session_state['mandant'].lower().replace(' ', '_')}.json"
    
    if os.path.exists(settings_file):
        try:
            with open(settings_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def save_settings(settings):
    """Speichert mandantenspezifische Settings"""
    if 'mandant' not in st.session_state:
        return
    
    settings_file = f"settings_{st.session_state['mandant'].lower().replace(' ', '_')}.json"
    
    with open(settings_file, 'w', encoding='utf-8') as f:
        json.dump(settings, f, indent=2, ensure_ascii=False)

def check_initial_setup():
    """Prüft ob Ersteinrichtung nötig - KORRIGIERTE VERSION"""
    # Prüfe zuerst ob wir schon in der Verarbeitung sind
    if any([
        'df_leit' in st.session_state and st.session_state.df_leit is not None,
        'processing_active' in st.session_state and st.session_state.processing_active,
        'datum_filter_confirmed' in st.session_state and st.session_state.datum_filter_confirmed
    ]):
        return True
    
    settings = load_settings()
    
    if not settings or 'current_config' not in settings:
        show_initial_setup()
        return False
    
    current_config_name = settings.get('current_config')
    if current_config_name and current_config_name in settings:
        config = settings[current_config_name]
        
        try:
            # OPTIMIERUNG: Einzelaufruf der neuen Hilfsfunktion
            apply_config_to_session(config)
        except Exception as e:
            st.error(f"❌ Fehler beim Laden der Konfiguration: {e}")
            show_initial_setup()
            return False
    
    return True

def show_initial_setup():
    """Zeigt Ersteinrichtungs-Dialog"""
    st.title("⚙️ Ersteinrichtung", help="Grundkonfiguration für Ihren Mandanten. Diese Werte gelten für alle zukünftigen Verarbeitungen.")
    st.info("Bitte konfigurieren Sie die Grundeinstellungen für Ihren Mandanten.")
    
    existing_settings = load_settings()
    has_existing = bool(existing_settings)
    
    with st.form("initial_setup"):
        st.subheader("Bewilligungszeitraum", help="Der genehmigte Zeitraum für Ihre Bürgschaft. Nur Vorgänge in diesem Zeitraum werden verarbeitet.")
        
        col1, col2 = st.columns(2)
        with col1:
            von_datum = st.date_input("Von", value=date.today() if not has_existing else date(2024, 5, 1),
                                    help="Beginn des Bewilligungszeitraums")
        with col2:
            default_bis = date.today().replace(year=date.today().year + 1) if not has_existing else date(2025, 4, 30)
            bis_datum = st.date_input("Bis", value=default_bis,
                                    help="Ende des Bewilligungszeitraums")
        
        st.subheader("Finanzielle Parameter", help="Grundlegende Werte für die Bürgschaftsberechnung")
        
        col3, col4 = st.columns(2)
        with col3:
            buergschaft = st.number_input(
                "Bürgschafts-Startsumme (€)",
                min_value=0.0,
                value=0.0,
                step=100000.0,
                format="%.2f",
                help="Die von Ihrer Bank/Versicherung bereitgestellte Bürgschaftssumme"
            )
            ersatz_zollsatz = st.number_input(
                "Ersatz-Zollsatz bei 0% (%)",
                min_value=0.0,
                max_value=100.0,
                value=0.0 if not has_existing else 12.0,
                step=0.1,
                help="Wird automatisch verwendet, wenn Waren mit 0% Zollsatz gefunden werden. Empfehlung: 12%"
            )
        
        with col4:
            pauschale = st.number_input(
                "Pauschalbetrag (€)",
                min_value=0.0,
                value=0.0 if not has_existing else 10000.0,
                step=1000.0,
                format="%.2f",
                help="Bürgschaftsbetrag für Positionen ohne ermittelbaren Zollwert. Verhindert 0€-Einträge in der Berechnung."
            )
        
        st.subheader("Bürgschaftserhöhung (optional)", help="Für unterjährige Erhöhungen der Bürgschaftssumme, z.B. bei Geschäftswachstum")
        buergschaft_erhoehung_aktiv = st.checkbox("Bürgschaftserhöhung aktivieren", value=False,
                                                 help="Aktivieren Sie diese Option, wenn die Bürgschaft während des Bewilligungszeitraums erhöht wurde")
        
        col5, col6 = st.columns(2)
        with col5:
            buergschaft_erhoehung_datum = st.date_input(
                "Datum der Erhöhung",
                value=date.today(),
                disabled=not buergschaft_erhoehung_aktiv,
                help="Ab wann gilt die erhöhte Bürgschaft?"
            )
        with col6:
            buergschaft_erhoehung_betrag = st.number_input(
                "Erhöhungsbetrag (€)",
                min_value=0.0,
                value=0.0,
                step=100000.0,
                format="%.2f",
                disabled=not buergschaft_erhoehung_aktiv,
                help="Um welchen Betrag wurde die Bürgschaft erhöht?"
            )
        
        if st.form_submit_button("💾 Einstellungen speichern", type="primary"):
            config_name = f"{von_datum.year}_{bis_datum.year}"
            
            settings = {
                config_name: {
                    "von": von_datum.strftime('%d.%m.%Y'),
                    "bis": bis_datum.strftime('%d.%m.%Y'),
                    "buergschaft": buergschaft,
                    "ersatz_zollsatz": ersatz_zollsatz,
                    "pauschale": pauschale,
                    "buergschaft_erhoehung_aktiv": buergschaft_erhoehung_aktiv,
                    "buergschaft_erhoehung_datum": buergschaft_erhoehung_datum.strftime('%d.%m.%Y'),
                    "buergschaft_erhoehung_betrag": buergschaft_erhoehung_betrag
                },
                "current_config": config_name
            }
            
            save_settings(settings)
            st.success("✅ Einstellungen gespeichert!")
            st.rerun()

def show_settings_page():
    """Zeigt Settings-Seite im Hauptbereich"""
    st.title("⚙️ Einstellungen")
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.button("← Zurück zur Verarbeitung", type="secondary"):
            st.session_state['show_settings'] = False
            st.rerun()
    
    settings = load_settings()
    current_config = settings.get(settings.get('current_config', ''), {})
    
    with st.form("settings_change"):
        st.subheader("Mandanten-Einstellungen bearbeiten", help="Ändern Sie hier die Grundkonfiguration für Ihre Bürgschaftsberechnung")
        
        st.markdown("### Bewilligungszeitraum", help="Der genehmigte Zeitraum für Ihre Bürgschaft")
        col1, col2 = st.columns(2)
        with col1:
            von_datum = st.date_input(
                "Von", 
                value=datetime.strptime(current_config.get('von', '01.05.2024'), '%d.%m.%Y').date(),
                help="Beginn des Bewilligungszeitraums"
            )
        with col2:
            bis_datum = st.date_input(
                "Bis", 
                value=datetime.strptime(current_config.get('bis', '30.04.2025'), '%d.%m.%Y').date(),
                help="Ende des Bewilligungszeitraums"
            )
        
        st.markdown("### Finanzielle Parameter", help="Grundlegende Werte für die Bürgschaftsberechnung")
        col3, col4 = st.columns(2)
        with col3:
            buergschaft = st.number_input(
                "Bürgschafts-Startsumme (€)",
                value=float(current_config.get('buergschaft', 0)),
                step=100000.0,
                format="%.2f",
                help="Die von Ihrer Bank/Versicherung bereitgestellte Bürgschaftssumme"
            )
            ersatz_zollsatz = st.number_input(
                "Ersatz-Zollsatz bei 0% (%)",
                value=float(current_config.get('ersatz_zollsatz', 0)),
                step=0.1,
                help="Wird automatisch verwendet, wenn Waren mit 0% Zollsatz gefunden werden. Empfehlung: 12%"
            )
        
        with col4:
            pauschale = st.number_input(
                "Pauschalbetrag (€)",
                value=float(current_config.get('pauschale', 0)),
                step=1000.0,
                format="%.2f",
                help="Bürgschaftsbetrag für Positionen ohne ermittelbaren Zollwert. Standard: 10.000€"
            )
        
        st.markdown("### Bürgschaftserhöhung (optional)", help="Für unterjährige Erhöhungen der Bürgschaftssumme")
        buergschaft_erhoehung_aktiv = st.checkbox(
            "Bürgschaftserhöhung aktivieren", 
            value=current_config.get('buergschaft_erhoehung_aktiv', False),
            help="Aktivieren Sie diese Option, wenn die Bürgschaft während des Bewilligungszeitraums erhöht wurde"
        )
        
        col5, col6 = st.columns(2)
        with col5:
            if current_config.get('buergschaft_erhoehung_datum'):
                datum_value = datetime.strptime(current_config['buergschaft_erhoehung_datum'], '%d.%m.%Y').date()
            else:
                datum_value = date.today()
                
            buergschaft_erhoehung_datum = st.date_input(
                "Datum der Erhöhung",
                value=datum_value,
                disabled=not buergschaft_erhoehung_aktiv,
                help="Ab wann gilt die erhöhte Bürgschaft?"
            )
        with col6:
            buergschaft_erhoehung_betrag = st.number_input(
                "Erhöhungsbetrag (€)",
                value=float(current_config.get('buergschaft_erhoehung_betrag', 0)),
                step=100000.0,
                format="%.2f",
                disabled=not buergschaft_erhoehung_aktiv,
                help="Um welchen Betrag wurde die Bürgschaft erhöht?"
            )
        
        submitted = st.form_submit_button("💾 Einstellungen speichern", type="primary", use_container_width=True)
        
    if submitted:
        config_name = f"{von_datum.year}_{bis_datum.year}"
        settings[config_name] = {
            "von": von_datum.strftime('%d.%m.%Y'),
            "bis": bis_datum.strftime('%d.%m.%Y'),
            "buergschaft": buergschaft,
            "ersatz_zollsatz": ersatz_zollsatz,
            "pauschale": pauschale,
            "buergschaft_erhoehung_aktiv": buergschaft_erhoehung_aktiv,
            "buergschaft_erhoehung_datum": buergschaft_erhoehung_datum.strftime('%d.%m.%Y'),
            "buergschaft_erhoehung_betrag": buergschaft_erhoehung_betrag
        }
        settings['current_config'] = config_name
        save_settings(settings)
        
        # OPTIMIERUNG: Einzelaufruf der neuen Hilfsfunktion
        apply_config_to_session(settings[config_name])
        
        st.success("✅ Einstellungen wurden gespeichert!")
        st.balloons()
        st.info("👆 Klicken Sie auf 'Zurück zur Verarbeitung' um fortzufahren.")

    st.markdown("---")
    st.subheader("📊 Erweiterte Informationen", help="Zusätzliche Verarbeitungsparameter und Einstellungen")
    
    tab1, tab2, tab3 = st.tabs(["Verarbeitungsregeln", "Import-Einstellungen", "Automatische Funktionen"])
    
    with tab1:
        col1, col2 = st.columns(2)
        with col1:
            st.info("**EUSt-Satz**")
            st.metric("Einfuhrumsatzsteuer", "19%")
            st.caption("Wird nur informativ berechnet, nicht Teil der Bürgschaft")
            
            st.info("**Verwahrungsfrist**")
            st.metric("Standard-Frist", "90 Tage")
            st.caption("Gestellungsdatum + 90 Tage")
        
        with col2:
            st.info("**WIDS-Verarbeitung**")
            st.write("Bei mehreren Positionen: **Position mit höchstem Zollwert**")
            st.caption("Pro Leitdatei-Zeile wird immer nur eine Ergebniszeile erstellt")
            
            st.info("**Leere Anmeldearten**")
            st.write("Verarbeitung: **Alle Zeilen (wie andere Anmeldearten)**")
            st.caption("Gilt für (leer), APDC, AVDC und NCAR")
    
    with tab2:
        st.info("**EZA-Spalten-Reduzierung**")
        st.write("✅ **Automatisch aktiviert**")
        st.caption("Reduziert große EZA-Dateien auf die 13 benötigten Standard-Spalten")
        
        st.info("**Zollsatz 0% Ersetzung**")
        st.write(f"✅ **Aktiviert** - Ersatz durch {st.session_state.get('zollsatz_ersatz', 0.12) * 100:.0f}%")
        st.caption("Für realistische Bürgschaftswerte bei zollfreien Waren")
        
        st.info("**BE-Anteil Verarbeitung**")
        st.write("✅ **Automatisch** - Zeilen werden nach BE-Anteil aufgeteilt")
        st.caption("Ermöglicht präzises 3-Kriterien-Matching für IMDC")
    
    with tab3:
        st.info("**Bürgschaftssaldo**")
        st.write("✅ **Wird immer berechnet**")
        st.caption("3 Excel-Sheets: Ergebnis, Bewegungsdetails, Tageszusammenfassung")
        
        st.info("**NCAR-Enhancement**")
        st.write("✅ **Immer aktiviert**")
        st.caption("Ergänzt automatisch Transport-MRN und Packstückzahlen")
        
        st.info("**ATB-Filter**")
        st.write("✅ **Automatisch** - S-Anmeldearten werden übersprungen")
        st.caption("Interne Konsolidierungen ohne Bürgschaftsrelevanz")
        
    st.info("""
    💡 **Hinweis:** Diese Einstellungen sind fest konfiguriert und gewährleisten eine konsistente Verarbeitung 
    gemäß den Zollvorschriften. Änderungen sind nur bei den Basis-Einstellungen möglich.
    """)

def show_downloads_section():
    """Zeigt die Downloads-Sektion"""
    st.title("📥 Downloads")
    
    if 'excel_file' not in st.session_state or st.session_state['excel_file'] is None:
        st.info("ℹ️ Noch keine Dokumente zum Download verfügbar.")
        st.write("Bitte führen Sie zuerst eine Verarbeitung durch:")
        st.write("1. Gehen Sie zum Tab '📊 Verarbeitung'")
        st.write("2. Laden Sie alle benötigten Dateien hoch")
        st.write("3. Starten Sie die Verarbeitung")
        st.write("")
        st.write("Nach erfolgreicher Verarbeitung stehen hier Ihre Dokumente zum Download bereit.")
        return
    
    st.success("✅ Ihre Dokumente sind bereit zum Download!")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Verarbeitete Zeilen", format_currency(len(st.session_state.get('ziel_sorted', [])), display_only=True))
    with col2:
        st.metric("Zeitraum", f"{st.session_state['von_datum'].strftime('%m/%Y')} - {st.session_state['bis_datum'].strftime('%m/%Y')}")
    with col3:
        st.metric("Max. Auslastung", st.session_state.get('max_auslastung', 'N/A'))
    
    st.markdown("---")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📊 Excel-Verwahrliste")
        
        mandant_prefix = st.session_state.get('mandant', 'Unbekannt')[:3].upper()
        von_str = st.session_state['von_datum'].strftime('%m_%y')
        bis_str = st.session_state['bis_datum'].strftime('%m_%y')
        excel_filename = f"Verwahrliste_{mandant_prefix}_{von_str}#{bis_str}.xlsx"
        
        st.download_button(
            label="📥 Excel herunterladen",
            data=st.session_state['excel_file'],
            file_name=excel_filename,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            help="Excel-Datei mit 3 Arbeitsblättern: Ergebnis, Bewegungsdetails, Tageszusammenfassung",
            key=f"download_excel_{datetime.now().timestamp()}"
        )
        
        st.info("""
        **Enthaltene Sheets:**
        - **Ergebnis**: Alle berechneten Positionen mit Tagessalden
        - **Bewegungsdetails**: Chronologische Ein-/Ausgänge
        - **Tageszusammenfassung**: Tägliche Salden mit Extremwerten
        """)
        
        if st.session_state.get('ncar_enabled', True) and 'df_ncar' in st.session_state and st.session_state['df_ncar'] is not None:
            st.success("✨ NCAR-Daten wurden automatisch ergänzt")
    
    with col2:
        st.subheader("📄 Zoll-Dokumentation")
        
        if st.button("📚 Dokumentation erstellen", use_container_width=True, type="primary"):
            with st.spinner("Dokumentation wird erstellt..."):
                doc_data = create_personalized_documentation()
                if doc_data:
                    doc_filename = f"Zoll_Dokumentation_{mandant_prefix}_{von_str}#{bis_str}.docx"
                    
                    st.session_state['doc_file'] = doc_data
                    st.session_state['doc_filename'] = doc_filename
                    st.success("✅ Dokumentation wurde erstellt!")
                    st.rerun()
                else:
                    st.error("❌ Dokumentation konnte nicht erstellt werden")
        
        if 'doc_file' in st.session_state and st.session_state['doc_file'] is not None:
            st.download_button(
                label="📥 Word-Dokument herunterladen",
                data=st.session_state['doc_file'],
                file_name=st.session_state.get('doc_filename', 'Zoll_Dokumentation.docx'),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
            st.info("""
            **Enthält:**
            - Prozessdokumentation
            - Verarbeitungsstatistiken
            - Konfigurationsübersicht
            - Hinweise für den Zoll
            """)
        
        st.warning("⚠️ Hinweis: Die Dokumentation benötigt die Datei 'Zoll_Dokumentation_Template.docx' im Programmverzeichnis.")

# === KONSTANTEN UND KONFIGURATION ===

VERARBEITBARE_ARTEN = ['IMDC', 'WIDS', 'IPDC', 'NCDP']
S_ANMELDEARTEN = ['SUSP', 'SUDC', 'SUCO', 'SUCF']
PAUSCHALE_ARTEN = ['(leer)', 'APDC', 'AVDC', 'NCAR']

EXAKTE_EZA_SPALTEN = [
    "Teilnehmer",
    "Verfahren", 
    "Bezugsnummer/LRN",
    "Überlassungsdatum",
    "Registriernummer/MRN",
    "PositionNo",
    "Zollwert",
    "AbgabeZoll",
    "AbgabeZollsatz",
    "Eustwert",
    "AbgabeEust",
    "Warentarifnummer",
    "BEAnteil SumA"
]

ANMELDEART_CONFIG = {
    'IMDC': {
        'unique_field': 'weitere',
        'import_source': 'df_import_eza',
        'match_field': 'import_eza_col',
        'pos_field': 'pos_field_eza',
        'has_menge': True,
        'needs_import': True,
        'process_all_rows': True
    },
    'WIDS': {
        'unique_field': 'weitere',
        'import_source': 'df_import_zl',
        'match_field': 'import_zl_col',
        'pos_field': 'pos_field_zl',
        'has_menge': False,
        'needs_import': True,
        'process_all_rows': True
    },
    'IPDC': {
        'unique_field': 'weitere',
        'needs_import': False,
        'process_all_rows': True
    },
    'NCDP': {
        'unique_field': 'reg',
        'import_source': 'df_ncts',
        'match_field': 'ncts_mrn_col',
        'needs_import': True,
        'process_all_rows': True
    },
    'NCAR': {
        'unique_field': 'reg',
        'needs_import': False,
        'process_all_rows': True
    }
}

DEFAULT_VALUES = {
    'df_leit': None,
    'df_import_eza': None,
    'df_import_zl': None,
    'df_ncts': None,
    'df_ncar': None,
    'stats': {},
    'leere_verarbeiten': True,
    'eust_satz': 0.19,
    'verwahrungsfrist_tage': 90,
    'wids_aggregation': 'Position mit höchstem Zollwert',
    'leere_anmeldeart_option': 'Alle Zeilen (wie andere Anmeldearten)',
    'zollsatz_null_ersetzen': True,
    'datum_filter_confirmed': False,
    'eza_auto_reduce': True,
    'ncar_enabled': True,
    'atb_filtered_count': 0,
    'show_settings': False,
    'processing_active': False,
    'processing_error': None
}

# === INITIALISIERUNG ===

def init_session_state():
    """Initialisiert den Session State"""
    for key, value in DEFAULT_VALUES.items():
        if key not in st.session_state:
            st.session_state[key] = value

# === UTILITY FUNKTIONEN ===

def apply_config_to_session(config: Dict):
    """NEUE HILFSFUNKTION: Lädt Konfigurationswerte sicher in den st.session_state."""
    st.session_state['von_datum'] = datetime.strptime(config.get('von', '01.05.2024'), '%d.%m.%Y').date()
    st.session_state['bis_datum'] = datetime.strptime(config.get('bis', '30.04.2025'), '%d.%m.%Y').date()
    st.session_state['startbuergschaft'] = float(config.get('buergschaft', 0))
    st.session_state['zollsatz_ersatz'] = float(config.get('ersatz_zollsatz', 12.0)) / 100
    st.session_state['pauschalbetrag'] = float(config.get('pauschale', 10000))
    st.session_state['buergschaft_erhöhung_aktiv'] = config.get('buergschaft_erhoehung_aktiv', False)
    erhoehung_datum_str = config.get('buergschaft_erhoehung_datum')
    if erhoehung_datum_str:
        st.session_state['buergschaft_erhöhung_datum'] = datetime.strptime(erhoehung_datum_str, '%d.%m.%Y').date()
    else:
        st.session_state['buergschaft_erhöhung_datum'] = date.today()
    st.session_state['buergschaft_erhöhung_betrag'] = float(config.get('buergschaft_erhoehung_betrag', 0))

def show_status(message, status='info', icon=None):
    """Einheitliche Status-Meldungen mit Icons"""
    icons = {
        'success': '✅',
        'info': 'ℹ️',
        'warning': '⚠️',
        'error': '❌'
    }
    
    display_icon = icon or icons.get(status, '')
    
    if status == 'success':
        st.success(f"{display_icon} {message}")
    elif status == 'warning':
        st.warning(f"{display_icon} {message}")
    elif status == 'error':
        st.error(f"{display_icon} {message}")
    else:
        st.info(f"{display_icon} {message}")

def format_currency(value, display_only=False, excel=False):
    """Formatiert Währung: Punkt als Tausender, Komma als Dezimal"""
    if display_only:
        return f"{int(value):,}".replace(',', '.')
    elif excel:
        return f"{value:,.2f}"
    else:
        formatted = f"{value:,.2f}"
        return formatted.replace(',', 'X').replace('.', ',').replace('X', '.')

def safe_strftime(dt) -> str:
    """Konvertiert ein Datum sicher in einen String."""
    try:
        return pd.to_datetime(dt).strftime('%d.%m.%Y') if pd.notnull(dt) else ''
    except (ValueError, TypeError):
        return ''

def safe_date_value(dt):
    """Konvertiert ein Datum sicher in ein Date-Objekt für Excel."""
    try:
        if pd.notnull(dt):
            parsed = pd.to_datetime(dt)
            return parsed.date() if hasattr(parsed, 'date') else parsed
        return None
    except (ValueError, TypeError):
        return None

def process_suma_position(row_data, suma_pos_col):
    """Verarbeitet SUMA-Position einheitlich"""
    if suma_pos_col and suma_pos_col in row_data:
        suma_value = row_data[suma_pos_col]
        return pd.to_numeric(suma_value, errors='ignore') if pd.notna(suma_value) else ''
    return ''

def prepare_dataframe_for_sorting(df):
    """Bereitet DataFrame für Standard-Sortierung vor"""
    df = df.copy()
    df['_gestell_date'] = df['Gestellungsdatum'].apply(parse_german_date)
    df['_suma_pos_numeric'] = pd.to_numeric(df['SUMA-Position'], errors='coerce').fillna(999999)
    return df

def sort_dataframe_standard(df):
    """Führt Standard-Sortierung durch und entfernt temporäre Spalten"""
    df_sorted = df.sort_values(
        by=['_gestell_date', 'ATB-Nummer', '_suma_pos_numeric'],
        ascending=[True, True, True]
    )
    columns_to_drop = ['_gestell_date', '_suma_pos_numeric']
    columns_to_drop = [col for col in columns_to_drop if col in df_sorted.columns]
    if columns_to_drop:
        df_sorted = df_sorted.drop(columns=columns_to_drop)
    return df_sorted

def find_col(df: pd.DataFrame, candidates: List[str], required: bool = True) -> Optional[str]:
    """OPTIMIERT: Findet die erste passende Spalte und stoppt bei kritischem Fehler."""
    for c in candidates:
        if c in df.columns:
            return c
    
    if required:
        st.error(f"❌ Kritischer Fehler: Eine der folgenden Pflichtspalten wurde in einer Ihrer Dateien nicht gefunden: `{', '.join(candidates)}`. Bitte prüfen Sie die Datei und laden Sie sie erneut hoch.")
        st.stop()
    return None

def clean_mrn(mrn_value) -> str:
    """Bereinigt MRN-Werte für den Vergleich"""
    if pd.isna(mrn_value):
        return ''
    cleaned = str(mrn_value).strip()
    if '.' in cleaned:
        cleaned = cleaned.split('.')[0]
    return cleaned

def has_atb_in_weitere_folge(leit_row, field_mappings) -> bool:
    """Prüft ob Weitere Registriernummer Folgeverfahren mit ATB beginnt"""
    weitere_folge = leit_row.get(field_mappings['leit_col_weitere'], '')
    return str(weitere_folge).strip().startswith('ATB')

def parse_german_date(date_str):
    """Konvertiert deutsches Datum (DD.MM.YYYY) in Date-Objekt"""
    if pd.isna(date_str) or date_str == '' or date_str is None:
        return None
    
    try:
        if isinstance(date_str, (datetime, date)):
            return date_str.date() if isinstance(date_str, datetime) else date_str
        
        if isinstance(date_str, str):
            date_str = date_str.strip()
            if '.' in date_str:
                parts = date_str.split('.')
                if len(parts) == 3:
                    try:
                        day, month, year = int(parts[0]), int(parts[1]), int(parts[2])
                        if 1 <= day <= 31 and 1 <= month <= 12 and 1900 <= year <= 2100:
                            return datetime(year, month, day).date()
                    except ValueError:
                        pass
            
            try:
                parsed = pd.to_datetime(date_str, format='%d.%m.%Y')
                return parsed.date()
            except Exception:
                pass
    except Exception:
        pass
    
    return None

def calculate_warehouse_dates(gestell, beendigung, frist_tage=90):
    """Zentrale Datumberechnung für alle Anmeldearten"""
    try:
        dt1 = pd.to_datetime(gestell)
        dt2 = pd.to_datetime(beendigung)
        verfrist_date = (dt1 + timedelta(days=frist_tage)).date()
        return {
            'verwahrungsfrist': (dt1 + timedelta(days=frist_tage)).strftime('%d.%m.%Y'),
            'verwahrungsfrist_date': verfrist_date,
            'verwahrungsdauer': (dt2 - dt1).days + 1
        }
    except Exception:
        return {'verwahrungsfrist': '', 'verwahrungsfrist_date': None, 'verwahrungsdauer': 0}

def validate_dataframe(df: pd.DataFrame, required_cols: List[List[str]], df_name: str) -> bool:
    """Validiert, ob alle erforderlichen Spalten vorhanden sind."""
    missing = []
    for col_candidates in required_cols:
        if not any(c in df.columns for c in col_candidates):
            missing.append(col_candidates)
    
    if missing:
        st.error(f"❌ Fehlende Spalten in {df_name}:")
        for cols in missing:
            st.error(f"   - Eine dieser Spalten wird benötigt: {cols}")
        raise ValueError(f"Pflichtfelder fehlen in {df_name}")
    return True

def validate_import_file(df, required_cols, file_name):
    """Einheitliche Validierung für Import-Dateien"""
    missing_cols = []
    for col in required_cols:
        if col not in df.columns:
            missing_cols.append(col)
    
    if missing_cols:
        st.error(f"❌ Fehlende Spalten in {file_name}: {missing_cols}")
        return False
    return True

def calculate_statistics(df: pd.DataFrame, anmeldeart_col: str) -> Dict[str, int]:
    """Berechnet Statistiken für die Anzeige."""
    stats = {}
    if anmeldeart_col in df.columns:
        alle_anmeldearten = VERARBEITBARE_ARTEN + S_ANMELDEARTEN + ['APDC', 'AVDC', 'NCAR']
        counts = df[anmeldeart_col].value_counts()
        
        for art in alle_anmeldearten:
            stats[f"{art}"] = int(counts.get(art, 0))
        
        stats["(leer)"] = int(df[anmeldeart_col].isna().sum())
        stats["Gesamt"] = len(df)
    return stats

def apply_zoelle_rule(results):
    """Wendet die Regel an: Mindestabgaben und Zölle = Gesamtabgaben"""
    pauschalbetrag = st.session_state.get('pauschalbetrag', 10000.0)
    
    for row in results:
        if row['Gesamtabgaben'] > 0 and row['Gesamtabgaben'] < 1.0:
            row['Gesamtabgaben'] = 1.0
        elif row['Gesamtabgaben'] == 0:
            if row['Zollwert (total)'] > 0:
                row['Gesamtabgaben'] = 1.0
            else:
                row['Gesamtabgaben'] = pauschalbetrag
        
        row['Zölle (total)'] = row['Gesamtabgaben']
    
    return results

def is_dataframe_valid(df):
    """Prüft ob DataFrame gültig und nicht leer ist"""
    return df is not None and isinstance(df, pd.DataFrame) and not df.empty

def clean_dataframe_for_export(df):
    """Bereinigt DataFrame von NaN-Werten für Excel-Export"""
    df_clean = df.copy()
    
    numeric_columns = ['Menge', 'Zollwert (total)', 'Drittlandzollsatz', 
                      'Zölle (total)', 'EUSt', 'Gesamtabgaben']
    for col in numeric_columns:
        if col in df_clean.columns:
            df_clean[col] = pd.to_numeric(df_clean[col], errors='coerce').fillna(0)
    
    text_columns = ['Referenznummer', 'MRN-Nummer Eingang', 'ATB-Nummer', 
                    'Erledigung mit', 'Anmeldeart']
    for col in text_columns:
        if col in df_clean.columns:
            df_clean[col] = df_clean[col].fillna('').astype(str)
    
    if 'SUMA-Position' in df_clean.columns:
        df_clean['SUMA-Position'] = df_clean['SUMA-Position'].apply(
            lambda x: pd.to_numeric(x, errors='ignore') if x != '' else ''
        )
    
    if 'Pos' in df_clean.columns:
        df_clean['Pos'] = df_clean['Pos'].apply(
            lambda x: pd.to_numeric(x, errors='ignore') 
            if x not in ['KEIN MATCH', 'Pauschale', ''] and not str(x).startswith('SUMME') and ' von ' not in str(x) else x
        )
    
    if 'Codenummer' in df_clean.columns:
        df_clean['Codenummer'] = df_clean['Codenummer'].apply(
            lambda x: pd.to_numeric(x, errors='ignore') if x != '' else ''
        )
    
    if 'Verwahrungsdauer' in df_clean.columns:
        df_clean['Verwahrungsdauer'] = pd.to_numeric(
            df_clean['Verwahrungsdauer'], errors='coerce'
        ).fillna(0).astype(int)
    
    return df_clean

def show_file_status(filename, count, reload_key, session_key=None, additional_info=None):
    """Zeigt Status einer geladenen Datei mit Reload-Option"""
    col1, col2 = st.columns([4, 1])
    with col1:
        st.success(f"✅ {filename} geladen: {count} Einträge")
        if additional_info:
            for info in additional_info:
                st.caption(info)
    with col2:
        if st.button("🔄", key=reload_key, help=f"{filename} neu laden"):
            if session_key:
                if session_key in st.session_state:
                    del st.session_state[session_key]
                if f"{session_key}_bytes" in st.session_state:
                    del st.session_state[f"{session_key}_bytes"]
            return True
    return False

# === GEMEINSAME VERARBEITUNGSFUNKTIONEN ===

def create_common_data(leit_row, gestell_col, dates_info):
    """Erstellt gemeinsame Daten für alle Anmeldearten"""
    return {
        'Referenznummer': str(leit_row.get('Bezugsnummer/LRN SumA', '')),
        'MRN-Nummer Eingang': str(leit_row.get('Registriernummer/MRN SumA', '')),
        'ATB-Nummer': str(leit_row.get('Registriernummer/MRN SumA', '')),
        'SUMA-Position': '',
        'Gestellungsdatum': safe_date_value(leit_row.get(gestell_col, '')),
        'Beendigung der Verwahrung': safe_date_value(leit_row.get('Datum Ende - CUSFIN', '')),
        'Verwahrungsfrist': dates_info.get('verwahrungsfrist_date', None),
        'Verwahrungsdauer': dates_info.get('verwahrungsdauer', 0),
        'Erledigung mit': ''
    }

def create_no_match_row(common_data, leit_row, anmeldeart, suma_pos_col):
    """Einheitliche No-Match Zeile für alle Anmeldearten"""
    common_data['SUMA-Position'] = process_suma_position(leit_row, suma_pos_col)
    
    pauschalbetrag = st.session_state.get('pauschalbetrag', 10000.0)
    
    if anmeldeart == 'NCDP':
        menge = 0
    elif anmeldeart == 'WIDS':
        menge = ''
    else:
        menge = 0
    
    return {
        **common_data,
        'Pos': 'KEIN MATCH',
        'Codenummer': '',
        'Menge': menge,
        'Zollwert (total)': 0.0,
        'Drittlandzollsatz': 0.0,
        'Zölle (total)': 0.0,
        'EUSt': 0.0,
        'Gesamtabgaben': pauschalbetrag,
        'Anmeldeart': anmeldeart
    }

def find_import_matches(unique_id, fallback_id, import_df, match_col):
    """Sucht Matches in Import-Datei mit Fallback"""
    matches = import_df[import_df[match_col] == unique_id]
    used_id = unique_id
    
    if matches.empty and fallback_id != unique_id:
        matches = import_df[import_df[match_col] == fallback_id]
        used_id = fallback_id
    
    return matches, used_id

def safe_numeric(value, default=0):
    """Konvertiert einen Wert sicher in eine Zahl"""
    result = pd.to_numeric(value, errors='coerce')
    return default if pd.isna(result) else float(result)

# === SPEZIFISCHE BERECHNUNGSFUNKTIONEN ===

def process_imdc_row(import_row, common_data, leit_row, pos_field, suma_pos_col):
    """Verarbeitet eine IMDC-Zeile"""
    zollwert = safe_numeric(import_row.get('Zollwert', 0))
    drittlandzollsatz = safe_numeric(import_row.get('AbgabeZollsatz', 0))
    
    if st.session_state.get('zollsatz_null_ersetzen', True) and drittlandzollsatz == 0 and zollwert > 0:
        drittlandzollsatz = st.session_state.get('zollsatz_ersatz', 0.12) * 100
    
    zoelle_total = round(zollwert * drittlandzollsatz / 100, 2)
    eust = round((zollwert + zoelle_total) * st.session_state.get('eust_satz', 0.19), 2)
    
    pauschalbetrag = st.session_state.get('pauschalbetrag', 10000.0)
    gesamtabgaben = zoelle_total if zollwert > 0 else pauschalbetrag
    
    common_data['ATB-Nummer'] = leit_row['Registriernummer/MRN SumA']
    common_data['SUMA-Position'] = process_suma_position(leit_row, suma_pos_col)
    
    menge_value = safe_numeric(import_row.get('Menge', 0))
    pos_value = import_row.get(pos_field, '')
    pos_numeric = pd.to_numeric(pos_value, errors='ignore') if pos_value else ''
    codenummer = import_row.get('Warentarifnummer', '')
    codenummer_numeric = pd.to_numeric(codenummer, errors='ignore') if codenummer else ''
    
    return {
        **common_data,
        'Pos': pos_numeric,
        'Codenummer': codenummer_numeric,
        'Menge': menge_value,
        'Zollwert (total)': zollwert,
        'Drittlandzollsatz': drittlandzollsatz,
        'Zölle (total)': zoelle_total,
        'EUSt': eust,
        'Gesamtabgaben': gesamtabgaben,
        'Anmeldeart': 'IMDC'
    }

def find_zl_value(row: pd.Series, field_type: str):
    """Findet Werte in der ZL-Datei mit verschiedenen Schreibweisen."""
    field_mapping = {
        'zollabgabe': ['Vorraussichtliche Zollabgabe', 'Voraussichtliche Zollabgabe'],
        'zollsatz': ['Vorraussichtliche Zollsatzabgabe', 'Voraussichtliche Zollsatzabgabe'],
        'dv1': ['DV1UmgerechnerterRechnungsbetrag', 'DV1 Umgerechneter Rechnungsbetrag']
    }
    
    candidates = field_mapping.get(field_type, [])
    for col in candidates:
        if col in row and pd.notna(row[col]):
            return safe_numeric(row[col])
    return 0.0

def calculate_wids_zollwert(row):
    """Berechnet den Zollwert für eine WIDS-Position"""
    zollabgabe = find_zl_value(row, 'zollabgabe')
    zollsatz = find_zl_value(row, 'zollsatz')
    dv1_betrag = find_zl_value(row, 'dv1')
    
    if zollsatz == 0:
        return dv1_betrag
    else:
        return round(zollabgabe / (zollsatz / 100), 2) if zollsatz > 0 else 0

def process_wids_row(import_row, common_data, leit_row, pos_field, suma_pos_col):
    """Verarbeitet eine WIDS-Zeile"""
    zollabgabe = find_zl_value(import_row, 'zollabgabe')
    zollsatz = find_zl_value(import_row, 'zollsatz')
    dv1_betrag = find_zl_value(import_row, 'dv1')
    
    if zollsatz == 0:
        zollwert = dv1_betrag
    else:
        zollwert = round(zollabgabe / (zollsatz / 100), 2) if zollsatz > 0 else 0
    
    if st.session_state.get('zollsatz_null_ersetzen', True) and zollsatz == 0 and zollwert > 0:
        zollsatz = st.session_state.get('zollsatz_ersatz', 0.12) * 100
        zollabgabe = round(zollwert * zollsatz / 100, 2)
    
    eust = round((zollwert + zollabgabe) * st.session_state.get('eust_satz', 0.19), 2)
    
    pauschalbetrag = st.session_state.get('pauschalbetrag', 10000.0)
    gesamtabgaben = zollabgabe if zollwert > 0 else pauschalbetrag
    
    common_data['SUMA-Position'] = process_suma_position(leit_row, suma_pos_col)
    
    pos_value = import_row.get(pos_field, '')
    pos_numeric = pd.to_numeric(pos_value, errors='ignore') if pos_value else ''
    codenummer = import_row.get('Warentarifnummer', '')
    codenummer_numeric = pd.to_numeric(codenummer, errors='ignore') if codenummer else ''
    
    return {
        **common_data,
        'Pos': pos_numeric,
        'Codenummer': codenummer_numeric,
        'Menge': '',
        'Zollwert (total)': zollwert,
        'Drittlandzollsatz': zollsatz,
        'Zölle (total)': zollabgabe,
        'EUSt': eust,
        'Gesamtabgaben': gesamtabgaben,
        'Anmeldeart': 'WIDS'
    }

def process_ipdc_row(common_data, leit_row, suma_pos_col):
    """Verarbeitet eine IPDC-Zeile"""
    zollwert_folge = safe_numeric(leit_row.get('Zollwert Folgeverfahren', 0))
    zollbetrag_folge = safe_numeric(leit_row.get('Zollbetrag Folgeverfahren', 0))
    
    if zollwert_folge > 0:
        drittlandzollsatz = round((zollbetrag_folge / zollwert_folge) * 100, 1)
    else:
        drittlandzollsatz = 0.0
    
    if st.session_state.get('zollsatz_null_ersetzen', True) and drittlandzollsatz == 0 and zollwert_folge > 0:
        drittlandzollsatz = st.session_state.get('zollsatz_ersatz', 0.12) * 100
        zollbetrag_folge = round(zollwert_folge * drittlandzollsatz / 100, 2)
    
    eust = round((zollwert_folge + zollbetrag_folge) * st.session_state.get('eust_satz', 0.19), 2)
    
    pauschalbetrag = st.session_state.get('pauschalbetrag', 10000.0)
    gesamtabgaben = zollbetrag_folge if zollwert_folge > 0 else pauschalbetrag
    
    if suma_pos_col and suma_pos_col in leit_row:
        suma_value = leit_row[suma_pos_col]
        common_data['SUMA-Position'] = pd.to_numeric(suma_value, errors='ignore') if pd.notna(suma_value) else ''
    else:
        common_data['SUMA-Position'] = ''
    
    common_data['Erledigung mit'] = common_data['Erledigung mit'] or str(leit_row.get('Weitere Registriernummer Folgeverfahren', ''))
    
    return {
        **common_data,
        'Pos': '',
        'Codenummer': '',
        'Menge': '',
        'Zollwert (total)': zollwert_folge,
        'Drittlandzollsatz': drittlandzollsatz,
        'Zölle (total)': zollbetrag_folge,
        'EUSt': eust,
        'Gesamtabgaben': gesamtabgaben,
        'Anmeldeart': 'IPDC'
    }

def extract_sicherheitsbetrag(sicherheit_data) -> float:
    """Extrahiert den Sicherheitsbetrag aus den NCTS-Sicherheitsdaten."""
    try:
        if isinstance(sicherheit_data, dict) and 'Sicherheitsleistungen' in sicherheit_data:
            leistungen = str(sicherheit_data['Sicherheitsleistungen'])
            match = re.search(r'Sicherheit:\s*([\d.,]+)', leistungen)
            if match:
                return float(match.group(1).replace(',', '.'))
        elif isinstance(sicherheit_data, str):
            match = re.search(r'Sicherheit:\s*([\d.,]+)', sicherheit_data)
            if match:
                return float(match.group(1).replace(',', '.'))
    except Exception as e:
        st.warning(f"⚠️ Fehler beim Extrahieren des Sicherheitsbetrags: {e}")
    return 0.0

def process_ncdp_row(common_data, leit_row, ncts_row, suma_pos_col):
    """Verarbeitet eine NCDP-Zeile"""
    sicherheitsbetrag = safe_numeric(extract_sicherheitsbetrag(ncts_row.get('Sicherheit', {})))
    
    common_data['SUMA-Position'] = process_suma_position(leit_row, suma_pos_col)
    
    return {
        **common_data,
        'Pos': '',
        'Codenummer': '',
        'Menge': 0,
        'Zollwert (total)': 0.0,
        'Drittlandzollsatz': 0.0,
        'Zölle (total)': 0.0,
        'EUSt': 0.0,
        'Gesamtabgaben': round(sicherheitsbetrag, 2),
        'Anmeldeart': 'NCDP'
    }

# === BE-ANTEIL VERARBEITUNG ===

def process_eza_be_anteil(df_import_eza):
    """Verarbeitet die BE Anteil SumA Spalte und multipliziert Zeilen entsprechend"""
    if len(df_import_eza.columns) < 13:
        st.warning("⚠️ Spalte 'BEAnteil SumA' nicht gefunden. Fahre ohne Verarbeitung fort.")
        return df_import_eza
    
    be_anteil_col = df_import_eza.columns[12]
    processed_rows = []
    
    for idx, row in df_import_eza.iterrows():
        be_anteil_value = row[be_anteil_col]
        
        if pd.isna(be_anteil_value) or str(be_anteil_value).strip() == '':
            new_row = row.copy()
            new_row['ATBnummer'] = ''
            new_row['Position'] = ''
            processed_rows.append(new_row)
        else:
            be_anteil_str = str(be_anteil_value).strip()
            atb_entries = be_anteil_str.split(',')
            
            for entry in atb_entries:
                entry = entry.strip()
                
                if ' - POS ' in entry:
                    try:
                        atb_part = entry.split(' - POS ')[0].strip()
                        pos_part = entry.split(' - POS ')[1].strip()
                        
                        new_row = row.copy()
                        new_row['ATBnummer'] = atb_part.strip()
                        new_row['Position'] = pos_part
                        processed_rows.append(new_row)
                    except Exception:
                        st.warning(f"⚠️ Konnte BE-Anteil nicht parsen: {entry}")
                        new_row = row.copy()
                        new_row['ATBnummer'] = ''
                        new_row['Position'] = ''
                        processed_rows.append(new_row)
                else:
                    new_row = row.copy()
                    new_row['ATBnummer'] = ''
                    new_row['Position'] = ''
                    processed_rows.append(new_row)
    
    return pd.DataFrame(processed_rows)

# === GENERISCHE ANMELDEARTEN-VERARBEITUNG ===

def process_anmeldeart_generic(anmeldeart, df_leit, data_sources, field_mappings, stats):
    """Generische Verarbeitung für alle Anmeldearten"""
    config = ANMELDEART_CONFIG.get(anmeldeart, {})
    results = []
    
    anmeldeart_data = df_leit[df_leit[field_mappings['anmeldeart_col']] == anmeldeart]
    
    for idx, leit_row in anmeldeart_data.iterrows():
        if has_atb_in_weitere_folge(leit_row, field_mappings):
            stats['atb_skipped'] = stats.get('atb_skipped', 0) + 1
            continue
        
        uid = leit_row[field_mappings[f'leit_col_{config["unique_field"]}']]
        
        dates_info = calculate_warehouse_dates(
            leit_row[field_mappings['gestell_col']], 
            leit_row['Datum Ende - CUSFIN'],
            st.session_state.get('verwahrungsfrist_tage', 90)
        )
        
        common_data = create_common_data(leit_row, field_mappings['gestell_col'], dates_info)
        
        results.extend(process_anmeldeart_row(
            anmeldeart, uid, leit_row, common_data, data_sources, field_mappings, stats
        ))
        
        stats[f'processed_{anmeldeart.lower()}'] += 1
    
    return results

def process_anmeldeart_row(anmeldeart, uid, leit_row, common_data, data_sources, field_mappings, stats):
    """Verarbeitet eine einzelne Zeile basierend auf der Anmeldeart"""
    if anmeldeart == 'IMDC':
        return process_imdc_generic(
            uid, leit_row, common_data, data_sources, field_mappings, stats
        )
    elif anmeldeart == 'WIDS':
        return process_wids_generic(
            uid, leit_row, common_data, data_sources, field_mappings, stats
        )
    elif anmeldeart == 'IPDC':
        zollwert = pd.to_numeric(leit_row.get('Zollwert Folgeverfahren', 0), errors='coerce')
        if zollwert > 0:
            stats['ipdc_with_zollwert'] += 1
        else:
            stats['ipdc_without_zollwert'] += 1
        return [process_ipdc_row(common_data, leit_row, field_mappings['suma_pos_col'])]
    elif anmeldeart == 'NCDP':
        return process_ncdp_generic(
            uid, leit_row, common_data, data_sources, field_mappings, stats
        )
    else:
        return []

def process_imdc_generic(uid, leit_row, common_data, data_sources, field_mappings, stats):
    """IMDC-spezifische Verarbeitung mit verbessertem 3-Kriterien-Matching"""
    results = []
    
    mrn_suma = leit_row.get('Registriernummer/MRN SumA', '')
    pos_suma = str(leit_row.get(field_mappings['suma_pos_col'], '')) if field_mappings['suma_pos_col'] else ''
    
    import_df = data_sources['df_import_eza']
    match_col = field_mappings['import_eza_col']
    
    has_be_anteil = 'ATBnummer' in import_df.columns and 'Position' in import_df.columns
    
    if mrn_suma and pos_suma and has_be_anteil:
        mrn_weitere = uid
        mrn_reg = leit_row[field_mappings['leit_col_reg']]
        
        precise_matches = import_df[
            (import_df[match_col] == mrn_weitere) &
            (import_df['ATBnummer'] == mrn_suma) &
            (import_df['Position'] == pos_suma)
        ]
        
        used_mrn = mrn_weitere
        
        if precise_matches.empty and mrn_reg != mrn_weitere:
            precise_matches = import_df[
                (import_df[match_col] == mrn_reg) &
                (import_df['ATBnummer'] == mrn_suma) &
                (import_df['Position'] == pos_suma)
            ]
            used_mrn = mrn_reg
        
        if not precise_matches.empty:
            common_data['Erledigung mit'] = used_mrn
            stats['imdc_match'] += 1
            stats['imdc_3criteria_match'] += 1
            if pd.notna(precise_matches.iloc[0].get('ATBnummer', '')):
                stats['imdc_be_anteil_rows'] += 1
            
            import_row = precise_matches.iloc[0]
            results.append(process_imdc_row(
                import_row, common_data, leit_row, 
                field_mappings['pos_field_eza'], 
                field_mappings['suma_pos_col']
            ))
            return results
    
    fallback_id = leit_row[field_mappings['leit_col_reg']]
    fallback_matches, used_id = find_import_matches(
        uid, fallback_id, import_df, match_col
    )
    
    if not fallback_matches.empty:
        common_data['Erledigung mit'] = used_id
        stats['imdc_match'] += 1
        stats['imdc_fallback_match'] += 1
        
        sorted_matches = fallback_matches.sort_values(by=field_mappings['pos_field_eza'])
        import_row = sorted_matches.iloc[0]
        
        if has_be_anteil and pd.notna(import_row.get('ATBnummer', '')):
            stats['imdc_be_anteil_rows'] += 1
        
        results.append(process_imdc_row(
            import_row, common_data, leit_row, 
            field_mappings['pos_field_eza'], 
            field_mappings['suma_pos_col']
        ))
    else:
        stats['imdc_no_match'] += 1
        results.append(create_no_match_row(common_data, leit_row, 'IMDC', field_mappings['suma_pos_col']))
    
    return results

def process_wids_generic(uid, leit_row, common_data, data_sources, field_mappings, stats):
    """WIDS-spezifische Verarbeitung mit konfigurierbarer Aggregation"""
    results = []
    fallback_id = leit_row[field_mappings['leit_col_reg']]
    
    import_matches, used_id = find_import_matches(
        uid, fallback_id,
        data_sources['df_import_zl'],
        field_mappings['import_zl_col']
    )
    
    common_data['Erledigung mit'] = used_id
    
    if import_matches.empty:
        stats['wids_no_match'] += 1
        results.append(create_no_match_row(common_data, leit_row, 'WIDS', field_mappings['suma_pos_col']))
    else:
        stats['wids_match'] += 1
        
        agg_mode = st.session_state.get('wids_aggregation', 'Position mit höchstem Zollwert')
        
        if len(import_matches) == 1:
            results.append(process_wids_row(
                import_matches.iloc[0], common_data, leit_row,
                field_mappings['pos_field_zl'],
                field_mappings['suma_pos_col']
            ))
        else:
            if agg_mode == "Nur Position 1":
                sorted_matches = import_matches.sort_values(by=field_mappings['pos_field_zl'])
                selected_row = sorted_matches.iloc[0]
                row_data = process_wids_row(
                    selected_row, common_data, leit_row,
                    field_mappings['pos_field_zl'],
                    field_mappings['suma_pos_col']
                )
                pos_value = row_data['Pos']
                row_data['Pos'] = f"{pos_value} (1 von {len(import_matches)})"
                
            elif agg_mode == "Position mit höchstem Zollwert":
                import_matches['_calculated_zollwert'] = import_matches.apply(
                    lambda row: calculate_wids_zollwert(row), axis=1
                )
                selected_row = import_matches.loc[import_matches['_calculated_zollwert'].idxmax()]
                row_data = process_wids_row(
                    selected_row, common_data, leit_row,
                    field_mappings['pos_field_zl'],
                    field_mappings['suma_pos_col']
                )
                pos_value = row_data['Pos']
                row_data['Pos'] = f"{pos_value} (max von {len(import_matches)})"
                
            else:
                total_zollabgabe = 0
                total_zollwert = 0
                max_zollsatz = 0
                
                for _, import_row in import_matches.iterrows():
                    zollabgabe = find_zl_value(import_row, 'zollabgabe')
                    zollsatz = find_zl_value(import_row, 'zollsatz')
                    dv1_betrag = find_zl_value(import_row, 'dv1')
                    
                    if zollsatz == 0:
                        zollwert = dv1_betrag
                    else:
                        zollwert = round(zollabgabe / (zollsatz / 100), 2) if zollsatz > 0 else 0
                    
                    total_zollabgabe += zollabgabe
                    total_zollwert += zollwert
                    max_zollsatz = max(max_zollsatz, zollsatz)
                
                avg_zollsatz = round(total_zollabgabe / total_zollwert * 100, 2) if total_zollwert > 0 else 0
                
                if st.session_state.get('zollsatz_null_ersetzen', True) and avg_zollsatz == 0 and total_zollwert > 0:
                    avg_zollsatz = st.session_state.get('zollsatz_ersatz', 0.12) * 100
                    total_zollabgabe = round(total_zollwert * avg_zollsatz / 100, 2)
                
                eust = round((total_zollwert + total_zollabgabe) * st.session_state.get('eust_satz', 0.19), 2)
                
                if field_mappings['suma_pos_col'] and field_mappings['suma_pos_col'] in leit_row:
                    suma_value = leit_row[field_mappings['suma_pos_col']]
                    common_data['SUMA-Position'] = pd.to_numeric(suma_value, errors='ignore') if pd.notna(suma_value) else ''
                
                pauschalbetrag = st.session_state.get('pauschalbetrag', 10000.0)
                
                row_data = common_data.copy()
                row_data.update({
                    'Pos': f'SUMME ({len(import_matches)} Pos.)',
                    'Codenummer': '',
                    'Menge': '',
                    'Zollwert (total)': total_zollwert,
                    'Drittlandzollsatz': avg_zollsatz,
                    'Zölle (total)': total_zollabgabe,
                    'EUSt': eust,
                    'Gesamtabgaben': total_zollabgabe if total_zollwert > 0 else pauschalbetrag,
                    'Anmeldeart': 'WIDS'
                })
            
            results.append(row_data)
    
    return results

def process_ncdp_generic(uid, leit_row, common_data, data_sources, field_mappings, stats):
    """NCDP-spezifische Verarbeitung"""
    results = []
    fallback_id = leit_row[field_mappings['leit_col_weitere']]
    
    ncts_matches, used_id = find_import_matches(
        uid, fallback_id,
        data_sources['df_ncts'],
        field_mappings['ncts_mrn_col']
    )
    
    common_data['Erledigung mit'] = used_id
    
    if ncts_matches.empty:
        stats['ncdp_no_match'] += 1
        results.append(create_no_match_row(common_data, leit_row, 'NCDP', field_mappings['suma_pos_col']))
    else:
        stats['ncdp_match'] += 1
        ncts_row = ncts_matches.iloc[0]
        results.append(process_ncdp_row(common_data, leit_row, ncts_row, field_mappings['suma_pos_col']))
    
    return results

def process_pauschale_anmeldeart(df_leit, field_mappings, stats, anmeldeart_filter=None, anmeldeart_name='(leer)'):
    """Verarbeitet pauschale Anmeldearten (leer, APDC, AVDC, NCAR)"""
    results = []
    
    if anmeldeart_filter is None:
        anmeldeart_data = df_leit[df_leit[field_mappings['anmeldeart_col']].isna() | (df_leit[field_mappings['anmeldeart_col']] == '')]
    else:
        anmeldeart_data = df_leit[df_leit[field_mappings['anmeldeart_col']] == anmeldeart_filter]
    
    pauschalbetrag = st.session_state.get('pauschalbetrag', 10000.0)
    
    for idx, pos_data in anmeldeart_data.iterrows():
        if has_atb_in_weitere_folge(pos_data, field_mappings):
            stats['atb_skipped'] = stats.get('atb_skipped', 0) + 1
            continue
        
        stats[f'{anmeldeart_name.lower()}_processed'] += 1
        
        dates_info = calculate_warehouse_dates(
            pos_data[field_mappings['gestell_col']],
            pos_data['Datum Ende - CUSFIN'],
            st.session_state.get('verwahrungsfrist_tage', 90)
        )
        
        pos_value = ''
        if field_mappings['suma_pos_col'] and field_mappings['suma_pos_col'] in pos_data:
            pos_raw = pos_data[field_mappings['suma_pos_col']]
            pos_value = pd.to_numeric(pos_raw, errors='ignore') if pd.notna(pos_raw) else ''
        
        results.append({
            'Referenznummer': str(pos_data.get('Bezugsnummer/LRN SumA', '')),
            'MRN-Nummer Eingang': str(pos_data.get('Registriernummer/MRN SumA', '')),
            'ATB-Nummer': str(pos_data.get('Registriernummer/MRN SumA', '')),
            'SUMA-Position': pos_value,
            'Gestellungsdatum': safe_date_value(pos_data[field_mappings['gestell_col']]),
            'Beendigung der Verwahrung': safe_date_value(pos_data['Datum Ende - CUSFIN']),
            'Verwahrungsfrist': dates_info.get('verwahrungsfrist_date', None),
            'Verwahrungsdauer': dates_info.get('verwahrungsdauer', 0),
            'Erledigung mit': '',
            'Pos': pos_value if pos_value else 'Pauschale',
            'Codenummer': '',
            'Menge': 0,
            'Zollwert (total)': 0.0,
            'Drittlandzollsatz': 0.0,
            'Zölle (total)': 0.0,
            'EUSt': 0.0,
            'Gesamtabgaben': pauschalbetrag,
            'Anmeldeart': anmeldeart_name
        })
    
    return results

# === BÜRGSCHAFTSSALDO FUNKTIONEN ===

def create_bewegungstabelle(df_ziel):
    """Erstellt eine Memory-Tabelle mit allen Bewegungen (Ein- und Ausgänge)"""
    bewegungen = []
    
    for idx, row in df_ziel.iterrows():
        gestell_date = parse_german_date(row['Gestellungsdatum'])
        if gestell_date:
            bewegungen.append({
                'Datum': gestell_date,
                'Datum_str': safe_strftime(row['Gestellungsdatum']),
                'Bewegungsart': 'Eingang',
                'ATB-Nummer': row['ATB-Nummer'],
                'Referenznummer': row['Referenznummer'],
                'Pos': row['Pos'],
                'SUMA-Position': row.get('SUMA-Position', ''),
                'Belastung': row['Gesamtabgaben'],
                'Entlastung': 0,
                'Anmeldeart': row['Anmeldeart'],
                '_original_idx': idx
            })
        
        beend_date = parse_german_date(row['Beendigung der Verwahrung'])
        if beend_date:
            bewegungen.append({
                'Datum': beend_date,
                'Datum_str': safe_strftime(row['Beendigung der Verwahrung']),
                'Bewegungsart': 'Ausgang',
                'ATB-Nummer': row['ATB-Nummer'],
                'Referenznummer': row['Referenznummer'],
                'Pos': row['Pos'],
                'SUMA-Position': row.get('SUMA-Position', ''),
                'Belastung': 0,
                'Entlastung': row['Gesamtabgaben'],
                'Anmeldeart': row['Anmeldeart'],
                '_original_idx': idx
            })
    
    for bewegung in bewegungen:
        suma_pos = bewegung.get('SUMA-Position', '')
        try:
            bewegung['_suma_pos_numeric'] = float(suma_pos) if suma_pos != '' else 999999
        except Exception:
            bewegung['_suma_pos_numeric'] = 999999
    
    bewegungen.sort(key=lambda x: (x['Datum'], 0 if x['Bewegungsart'] == 'Eingang' else 1, x['_original_idx'], x['_suma_pos_numeric']))
    
    for bewegung in bewegungen:
        bewegung.pop('_suma_pos_numeric', None)
    
    df_bewegungen = pd.DataFrame(bewegungen)
    if '_original_idx' in df_bewegungen.columns:
        df_bewegungen = df_bewegungen.drop(columns=['_original_idx'])
    
    return df_bewegungen

def calculate_daily_summary(bewegungen_df, startbuergschaft):
    """Berechnet Tagessummen und fortlaufenden Bürgschaftsstand"""
    startbuergschaft = float(startbuergschaft)
    daily_summary = {}
    
    bewegungen_df['Belastung'] = pd.to_numeric(bewegungen_df['Belastung'], errors='coerce').fillna(0)
    bewegungen_df['Entlastung'] = pd.to_numeric(bewegungen_df['Entlastung'], errors='coerce').fillna(0)
    
    unique_dates = sorted(bewegungen_df['Datum'].unique())
    
    for datum in unique_dates:
        tages_data = bewegungen_df[bewegungen_df['Datum'] == datum]
        
        belastung_summe = float(tages_data['Belastung'].sum())
        entlastung_summe = float(tages_data['Entlastung'].sum())
        
        if (st.session_state.get('buergschaft_erhöhung_aktiv', False) and 
            datum == st.session_state.get('buergschaft_erhöhung_datum', date(2025, 2, 4))):
            entlastung_summe += st.session_state.get('buergschaft_erhöhung_betrag', 1500000.0)
        
        daily_summary[datum] = {
            'Belastung': round(belastung_summe, 2),
            'Entlastung': round(entlastung_summe, 2),
            'Netto': round(entlastung_summe - belastung_summe, 2)
        }
    
    laufender_stand = startbuergschaft
    for datum in unique_dates:
        laufender_stand = laufender_stand - daily_summary[datum]['Belastung'] + daily_summary[datum]['Entlastung']
        daily_summary[datum]['Bürgschaftsstand'] = round(laufender_stand, 2)
    
    return daily_summary

def add_tagessummen_to_ziel(df_ziel, daily_summary):
    """Fügt Tagessummen zur Zieldatei hinzu - in der letzten Zeile des Tages"""
    df_ziel = prepare_dataframe_for_sorting(df_ziel)
    df_sorted = df_ziel.sort_values(['_gestell_date', 'ATB-Nummer', '_suma_pos_numeric']).copy()
    
    df_sorted[''] = ''
    df_sorted['Belastung'] = ''
    df_sorted['Entlastung'] = ''
    df_sorted['Netto-Belastung'] = ''
    df_sorted['Bürgschaftsstand'] = ''
    
    for datum in daily_summary.keys():
        mask = df_sorted['_gestell_date'] == datum
        tag_indices = df_sorted[mask].index
        
        if len(tag_indices) > 0:
            last_idx = tag_indices[-1]
            
            if (st.session_state.get('buergschaft_erhöhung_aktiv', False) and 
                datum == st.session_state.get('buergschaft_erhöhung_datum', date(2025, 2, 4))):
                betrag = st.session_state.get('buergschaft_erhöhung_betrag', 1500000.0)
                df_sorted.loc[last_idx, ''] = f'TAGESSALDO {datum.strftime("%d.%m.%Y")} (Bürgschaft +{betrag/1000000:.1f} Mio)'
            else:
                df_sorted.loc[last_idx, ''] = f'TAGESSALDO {datum.strftime("%d.%m.%Y")}'
            
            df_sorted.loc[last_idx, 'Belastung'] = daily_summary[datum]['Belastung']
            df_sorted.loc[last_idx, 'Entlastung'] = daily_summary[datum]['Entlastung']
            df_sorted.loc[last_idx, 'Netto-Belastung'] = daily_summary[datum]['Netto']
            df_sorted.loc[last_idx, 'Bürgschaftsstand'] = daily_summary[datum]['Bürgschaftsstand']
    
    columns_to_drop = ['_gestell_date', '_suma_pos_numeric']
    columns_to_drop = [col for col in columns_to_drop if col in df_sorted.columns]
    if columns_to_drop:
        df_sorted = df_sorted.drop(columns=columns_to_drop)
    
    return df_sorted

def create_bewegungsdetails_df(bewegungen_df, daily_summary, startbuergschaft):
    """Erstellt die Bewegungsdetails-Tabelle mit Tagessummen"""
    bewegungen_df['_suma_pos_numeric'] = bewegungen_df['SUMA-Position'].apply(
        lambda x: float(x) if isinstance(x, (int, float)) or (isinstance(x, str) and x.replace('.','').isdigit()) else 999999
    )
    
    bewegungen_sorted = bewegungen_df.sort_values(['Datum', 'Bewegungsart', 'ATB-Nummer', '_suma_pos_numeric'], 
                                                  ascending=[True, False, True, True]).copy()
    bewegungen_sorted = bewegungen_sorted.drop(columns=['_suma_pos_numeric'])
    
    result_rows = []
    
    result_rows.append({
        'Datum': None,
        'ATB-Nummer': 'START',
        'Referenznummer': '',
        'SUMA-Position': '',
        'Pos': '',
        'Belastung': 0,
        'Entlastung': 0,
        'Netto-Belastung': 0,
        'Bürgschaftsstand': startbuergschaft
    })
    
    laufender_stand = float(startbuergschaft)
    current_date = None
    
    for idx, row in bewegungen_sorted.iterrows():
        datum_obj = row['Datum']
        
        if (current_date != datum_obj and 
            datum_obj == st.session_state.get('buergschaft_erhöhung_datum', date(2025, 2, 4)) and
            st.session_state.get('buergschaft_erhöhung_aktiv', False)):
            
            laufender_stand += st.session_state.get('buergschaft_erhöhung_betrag', 1500000.0)
            
            result_rows.append({
                'Datum': datum_obj,
                'ATB-Nummer': 'BÜRGSCHAFTSERHÖHUNG',
                'Referenznummer': 'Erhöhung der verfügbaren Bürgschaft',
                'SUMA-Position': '',
                'Pos': '',
                'Belastung': '',
                'Entlastung': st.session_state.get('buergschaft_erhöhung_betrag', 1500000.0),
                'Netto-Belastung': '',
                'Bürgschaftsstand': round(laufender_stand, 2)
            })
        
        if current_date is not None and datum_obj != current_date:
            if current_date in daily_summary:
                result_rows.append({
                    'Datum': current_date,
                    'ATB-Nummer': 'TAGESSUMME',
                    'Referenznummer': '',
                    'SUMA-Position': '',
                    'Pos': '',
                    'Belastung': daily_summary[current_date]['Belastung'],
                    'Entlastung': daily_summary[current_date]['Entlastung'],
                    'Netto-Belastung': daily_summary[current_date]['Belastung'] - daily_summary[current_date]['Entlastung'],
                    'Bürgschaftsstand': daily_summary[current_date]['Bürgschaftsstand']
                })
                
                result_rows.append({
                    'Datum': None,
                    'ATB-Nummer': '',
                    'Referenznummer': '',
                    'SUMA-Position': '',
                    'Pos': '',
                    'Belastung': '',
                    'Entlastung': '',
                    'Netto-Belastung': '',
                    'Bürgschaftsstand': ''
                })
        
        belastung_wert = float(row['Belastung'])
        entlastung_wert = float(row['Entlastung'])
        
        laufender_stand = laufender_stand - belastung_wert + entlastung_wert
        
        result_rows.append({
            'Datum': datum_obj,
            'ATB-Nummer': row['ATB-Nummer'],
            'Referenznummer': row['Referenznummer'],
            'SUMA-Position': row.get('SUMA-Position', ''),
            'Pos': row['Pos'] if pd.notna(row['Pos']) else '',
            'Belastung': belastung_wert if belastung_wert > 0 else '',
            'Entlastung': entlastung_wert if entlastung_wert > 0 else '',
            'Netto-Belastung': '',
            'Bürgschaftsstand': round(laufender_stand, 2)
        })
        
        current_date = datum_obj
    
    if current_date is not None and current_date in daily_summary:
        result_rows.append({
            'Datum': current_date,
            'ATB-Nummer': 'TAGESSUMME',
            'Referenznummer': '',
            'SUMA-Position': '',
            'Pos': '',
            'Belastung': daily_summary[current_date]['Belastung'],
            'Entlastung': daily_summary[current_date]['Entlastung'],
            'Netto-Belastung': daily_summary[current_date]['Belastung'] - daily_summary[current_date]['Entlastung'],
            'Bürgschaftsstand': daily_summary[current_date]['Bürgschaftsstand']
        })
    
    return pd.DataFrame(result_rows)

def create_tageszusammenfassung_df_mit_extrema(bewegungen_df, daily_summary, startbuergschaft):
    """Erstellt eine kompakte Tagesübersicht mit Tagessummen und Höchst-/Tiefstständen"""
    result_rows = []
    
    result_rows.append({
        'Datum': 'START',
        'Tages-Belastung': '',
        'Tages-Entlastung': '',
        'Netto-Bewegung': '',
        'Tiefststand': float(startbuergschaft),
        'Höchststand': float(startbuergschaft),
        'Schlussstand': float(startbuergschaft),
        'Auslastung %': 0.0,
        'Hinweis': ''
    })
    
    bewegungen_df['_suma_pos_numeric'] = bewegungen_df['SUMA-Position'].apply(
        lambda x: float(x) if isinstance(x, (int, float)) or (isinstance(x, str) and x.replace('.','').isdigit()) else 999999
    )
    
    bewegungen_sorted = bewegungen_df.sort_values(['Datum', 'Bewegungsart', 'ATB-Nummer', '_suma_pos_numeric'], 
                                                  ascending=[True, False, True, True]).copy()
    bewegungen_sorted = bewegungen_sorted.drop(columns=['_suma_pos_numeric'])
    
    laufender_stand = float(startbuergschaft)
    sorted_dates = sorted(daily_summary.keys())
    
    for datum in sorted_dates:
        tages_bewegungen = bewegungen_sorted[bewegungen_sorted['Datum'] == datum]
        tages_data = daily_summary[datum]
        
        tagesstart_stand = laufender_stand
        tiefststand = tagesstart_stand
        hoechststand = tagesstart_stand
        
        for _, bewegung in tages_bewegungen.iterrows():
            laufender_stand = laufender_stand - float(bewegung['Belastung']) + float(bewegung['Entlastung'])
            tiefststand = min(tiefststand, laufender_stand)
            hoechststand = max(hoechststand, laufender_stand)
        
        if (st.session_state.get('buergschaft_erhöhung_aktiv', False) and 
            datum == st.session_state.get('buergschaft_erhöhung_datum', date(2025, 2, 4))):
            laufender_stand += st.session_state.get('buergschaft_erhöhung_betrag', 1500000.0)
            hoechststand = max(hoechststand, laufender_stand)
        
        max_auslastung = 0 if startbuergschaft == 0 else ((startbuergschaft - tiefststand) / startbuergschaft * 100)
        
        hinweis = ''
        if (st.session_state.get('buergschaft_erhöhung_aktiv', False) and 
            datum == st.session_state.get('buergschaft_erhöhung_datum', date(2025, 2, 4))):
            betrag = st.session_state.get('buergschaft_erhöhung_betrag', 1500000.0)
            hinweis = f'Bürgschaftserhöhung +{betrag:,.0f} €'
        
        result_rows.append({
            'Datum': safe_strftime(datum),
            'Tages-Belastung': tages_data['Belastung'],
            'Tages-Entlastung': tages_data['Entlastung'],
            'Netto-Bewegung': tages_data['Belastung'] - tages_data['Entlastung'],
            'Tiefststand': round(tiefststand, 2),
            'Höchststand': round(hoechststand, 2),
            'Schlussstand': round(tages_data['Bürgschaftsstand'], 2),
            'Auslastung %': round(max_auslastung, 2),
            'Hinweis': hinweis
        })
    
    if len(daily_summary) > 0:
        total_belastung = sum(daily_summary[d]['Belastung'] for d in daily_summary)
        total_entlastung = sum(daily_summary[d]['Entlastung'] for d in daily_summary)
        final_stand = daily_summary[sorted_dates[-1]]['Bürgschaftsstand']
        
        alle_tiefstwerte = [row['Tiefststand'] for row in result_rows[1:] if isinstance(row['Tiefststand'], (int, float))]
        alle_hoechstwerte = [row['Höchststand'] for row in result_rows[1:] if isinstance(row['Höchststand'], (int, float))]
        
        globaler_tiefststand = min(alle_tiefstwerte) if alle_tiefstwerte else startbuergschaft
        globaler_hoechststand = max(alle_hoechstwerte) if alle_hoechstwerte else startbuergschaft
        max_auslastung = 0 if startbuergschaft == 0 else ((startbuergschaft - globaler_tiefststand) / startbuergschaft * 100)
        
        result_rows.append({
            'Datum': '',
            'Tages-Belastung': '',
            'Tages-Entlastung': '',
            'Netto-Bewegung': '',
            'Tiefststand': '',
            'Höchststand': '',
            'Schlussstand': '',
            'Auslastung %': '',
            'Hinweis': ''
        })
        
        result_rows.append({
            'Datum': 'GESAMT',
            'Tages-Belastung': round(total_belastung, 2),
            'Tages-Entlastung': round(total_entlastung, 2),
            'Netto-Bewegung': round(total_belastung - total_entlastung, 2),
            'Tiefststand': globaler_tiefststand,
            'Höchststand': globaler_hoechststand,
            'Schlussstand': final_stand,
            'Auslastung %': round(max_auslastung, 2),
            'Hinweis': ''
        })
    
    return pd.DataFrame(result_rows)

# === NCAR-ENHANCEMENT FUNKTIONEN ===

def enhance_ziel_with_ncar(ziel_df, ncar_df):
    """Erweitert Zieldatei um NCAR-Daten - vereinfachte Logik"""
    ziel_df['_atb_clean'] = ziel_df['ATB-Nummer'].astype(str).str.strip()
    ncar_df['_atb_clean'] = ncar_df['Registriernr.-SumA'].astype(str).str.strip()
    
    enhanced = ziel_df.merge(
        ncar_df[['_atb_clean', 'RegistriernNr./MRN', 'Anzahl Packstücke']],
        on='_atb_clean',
        how='left'
    )
    
    enhanced['MRN-Nummer Eingang'] = enhanced['RegistriernNr./MRN'].fillna(enhanced['MRN-Nummer Eingang'])
    enhanced['Menge'] = enhanced['Anzahl Packstücke'].fillna(enhanced['Menge'])
    
    enhanced = enhanced.drop(columns=['_atb_clean', 'RegistriernNr./MRN', 'Anzahl Packstücke'])
    
    return enhanced

def process_ncar_file(ncar_file):
    """Verarbeitet NCAR-Datei und validiert Struktur"""
    try:
        ncar_df = pd.read_excel(ncar_file)
        
        required_cols = ['Registriernr.-SumA', 'RegistriernNr./MRN', 'Anzahl Packstücke']
        
        if not validate_import_file(ncar_df, required_cols, "NCAR-Datei"):
            return None
        
        return ncar_df
        
    except Exception as e:
        st.error(f"❌ Fehler beim Lesen der NCAR-Datei: {e}")
        return None

def create_personalized_documentation():
    """Erstellt personalisierte Dokumentation basierend auf aktuellen Daten"""
    from docx import Document
    from docx.shared import Inches
    import io
    import os
    
    try:
        doc = Document("Zoll_Dokumentation_Template.docx")
        
        logo_path = st.session_state.get('mandant_logo')
        
        values = {
            "{{DATEINAME}}": f"Verwahrliste {st.session_state['mandant'][:3].upper()}_{st.session_state['von_datum'].strftime('%m_%y')}#{st.session_state['bis_datum'].strftime('%m_%y')}",
            "{{MANDANT}}": st.session_state.get('mandant', 'Unbekannt'),
            "{{HAUPTZOLLAMT}}": "Krefeld",
            "{{VON_DATUM}}": st.session_state['von_datum'].strftime('%d.%m.%Y'),
            "{{BIS_DATUM}}": st.session_state['bis_datum'].strftime('%d.%m.%Y'),
            "{{STARTBUERGSCHAFT}}": format_currency(st.session_state.get('startbuergschaft', 0)),
            "{{LEITDATEI_GESAMT}}": format_currency(st.session_state.stats.get('Gesamt', 0), display_only=True),
            "{{ERGEBNIS_ZEILEN}}": format_currency(len(st.session_state.get('ziel_sorted', [])), display_only=True),
            
            "{{S_ARTEN_SUMME}}": format_currency(sum(st.session_state.stats.get(art, 0) for art in S_ANMELDEARTEN), display_only=True),
            
            "{{PAUSCHALBETRAG}}": f"{st.session_state.get('pauschalbetrag', 10000):,.0f} €".replace(',', '.'),
            "{{ERSATZ_ZOLLSATZ}}": f"{st.session_state.get('zollsatz_ersatz', 0.12) * 100:.0f}".replace('.', ','),
            "{{EUST_SATZ}}": "19",
            "{{VERWAHRUNGSFRIST}}": "90",
            
            "{{MAX_AUSLASTUNG}}": st.session_state.get('max_auslastung_str', 'N/A'),
            "{{TIEFSTSTAND}}": st.session_state.get('tiefststand_str', 'N/A')
        }
        
        if st.session_state.get('buergschaft_erhöhung_aktiv', False):
            betrag = st.session_state.get('buergschaft_erhöhung_betrag', 0)
            datum = st.session_state.get('buergschaft_erhöhung_datum', date.today())
            values["{{BUERGSCHAFT_ERHOEHUNG}}"] = f", erhöht um {format_currency(betrag)} am {datum.strftime('%d.%m.%Y')}"
        else:
            values["{{BUERGSCHAFT_ERHOEHUNG}}"] = ""
        
        s_details = []
        for art in S_ANMELDEARTEN:
            if st.session_state.stats.get(art, 0) > 0:
                s_details.append(f"{art}: {format_currency(st.session_state.stats.get(art, 0), display_only=True)}")
        values["{{S_ARTEN_DETAILS}}"] = ", ".join(s_details) if s_details else "keine"
        
        for art in ['IMDC', 'IPDC', 'WIDS', 'NCDP', 'APDC', 'AVDC', 'NCAR']:
            count = st.session_state.stats.get(art, 0)
            values[f"{{{{{art}_COUNT}}}}"] = format_currency(count, display_only=True)
        
        values["{{LEER_COUNT}}"] = format_currency(st.session_state.stats.get('(leer)', 0), display_only=True)
        
        for art in S_ANMELDEARTEN:
            values[f"{{{{{art}_COUNT}}}}"] = format_currency(st.session_state.stats.get(art, 0), display_only=True)
        
        import_files = {
            'IMPORT_EZA': 'df_import_eza',
            'IMPORT_ZL': 'df_import_zl', 
            'NCTS': 'df_ncts',
            'NCAR': 'df_ncar'
        }
        
        for key, df_name in import_files.items():
            df = st.session_state.get(df_name)
            count = len(df) if df is not None else 0
            values[f"{{{{{key}_COUNT}}}}"] = format_currency(count, display_only=True)
        
        if 'ziel_sorted' in st.session_state and st.session_state['ziel_sorted'] is not None:
            ziel = st.session_state['ziel_sorted']
            
            ziel_numeric = ziel.copy()
            ziel_numeric['Zollwert (total)'] = pd.to_numeric(ziel_numeric['Zollwert (total)'], errors='coerce').fillna(0)
            
            values["{{GESAMT_ZOLLWERT}}"] = format_currency(ziel_numeric['Zollwert (total)'].sum())
            values["{{GESAMT_ZOELLE}}"] = format_currency(ziel['Zölle (total)'].sum())
            values["{{GESAMT_EUST}}"] = format_currency(ziel['EUSt'].sum())
            values["{{GESAMT_ABGABEN}}"] = format_currency(ziel['Gesamtabgaben'].sum())
            
            anmeldearten = ['IMDC', 'IPDC', 'WIDS', 'NCDP', '(leer)', 'APDC', 'AVDC', 'NCAR']
            
            for art in anmeldearten:
                art_data = ziel[ziel['Anmeldeart'] == art]
                
                if art == '(leer)':
                    art_key = 'LEER'
                else:
                    art_key = art
                
                if len(art_data) > 0:
                    art_numeric = art_data.copy()
                    art_numeric['Zollwert (total)'] = pd.to_numeric(art_numeric['Zollwert (total)'], errors='coerce').fillna(0)
                    
                    values[f"{{{{{art_key}_ZOLLWERT}}}}"] = format_currency(art_numeric['Zollwert (total)'].sum())
                    values[f"{{{{{art_key}_ZOELLE}}}}"] = format_currency(art_data['Zölle (total)'].sum())
                    values[f"{{{{{art_key}_EUST}}}}"] = format_currency(art_data['EUSt'].sum())
                    values[f"{{{{{art_key}_ABGABEN}}}}"] = format_currency(art_data['Gesamtabgaben'].sum())
                else:
                    values[f"{{{{{art_key}_ZOLLWERT}}}}"] = "0,00 €"
                    values[f"{{{{{art_key}_ZOELLE}}}}"] = "0,00 €"
                    values[f"{{{{{art_key}_EUST}}}}"] = "0,00 €"
                    values[f"{{{{{art_key}_ABGABEN}}}}"] = "0,00 €"
        else:
            values["{{GESAMT_ZOLLWERT}}"] = "0,00 €"
            values["{{GESAMT_ZOELLE}}"] = "0,00 €"
            values["{{GESAMT_EUST}}"] = "0,00 €"
            values["{{GESAMT_ABGABEN}}"] = "0,00 €"
            
            for art in ['IMDC', 'IPDC', 'WIDS', 'NCDP', 'LEER', 'APDC', 'AVDC', 'NCAR']:
                values[f"{{{{{art}_ZOLLWERT}}}}"] = "0,00 €"
                values[f"{{{{{art}_ZOELLE}}}}"] = "0,00 €"
                values[f"{{{{{art}_EUST}}}}"] = "0,00 €"
                values[f"{{{{{art}_ABGABEN}}}}"] = "0,00 €"
        
        for paragraph in doc.paragraphs:
            for key, value in values.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value))
        
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                if '{{LOGO}}' in paragraph.text:
                    paragraph.text = ''
                    if logo_path and os.path.exists(logo_path):
                        try:
                            run = paragraph.add_run()
                            run.add_picture(logo_path, width=Inches(1.5))
                        except Exception as e:
                            st.warning(f"Logo konnte nicht in Kopfzeile eingefügt werden: {e}")
                else:
                    for key, value in values.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in values.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, str(value))
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output.getvalue()
        
    except Exception as e:
        st.error(f"Fehler bei Dokumentenerstellung: {e}")
        return None

# === HAUPTVERARBEITUNG ===

def update_progress(progress_bar, current, total, prefix="", suffix=""):
    """Aktualisiert die Progress Bar mit Text"""
    percentage = current / total
    if prefix and suffix:
        text = f"{prefix} {suffix}"
    elif prefix:
        text = prefix
    elif suffix:
        text = suffix
    else:
        text = f"{int(percentage * 100)}%"
    
    progress_bar.progress(percentage, text=text)

def process_data():
    """Hauptverarbeitungsfunktion mit ATB-Filter und Error Handling"""
    st.session_state['processing_active'] = True
    st.session_state['processing_error'] = None
    
    with st.expander("📋 Verarbeitungsdetails anzeigen", expanded=False):
        schritte_container = st.container()
    
    with st.spinner("🔄 Verarbeitung wird vorbereitet..."):
        time.sleep(0.5)
    
    try:
        df_leit = st.session_state.df_leit.copy()
        
        data_sources = {
            'df_leit': df_leit,
            'df_import_eza': st.session_state.df_import_eza.copy() if st.session_state.df_import_eza is not None else pd.DataFrame(),
            'df_import_zl': st.session_state.df_import_zl.copy() if st.session_state.df_import_zl is not None else pd.DataFrame(),
            'df_ncts': st.session_state.df_ncts.copy() if is_dataframe_valid(st.session_state.get('df_ncts')) else pd.DataFrame()
        }
        
        progress_bar = st.progress(0, text="📊 Initialisiere Datenverarbeitung...")
        
        with schritte_container:
            st.success("✅ Datenverarbeitung initialisiert")
        
        # OPTIMIERUNG: Der try-except-Block wird durch die verbesserte find_col Funktion überflüssig.
        field_mappings = {
            'leit_col_weitere': find_col(df_leit, ['Weitere Registriernummer Folgeverfahren', 'Weitere Registriernummer']),
            'leit_col_reg': find_col(df_leit, ['Registriernummer Folgeverfahren']),
            'anmeldeart_col': find_col(df_leit, ['Anmeldeart Folgeverfahren']),
            'gestell_col': find_col(df_leit, ['Datum Überlassung - CUSTST']),
            'import_eza_col': find_col(data_sources['df_import_eza'], ['Registriernummer/MRN', 'Registriernummer / MRN', 'MRN'], required=not data_sources['df_import_eza'].empty) if not data_sources['df_import_eza'].empty else None,
            'import_zl_col': find_col(data_sources['df_import_zl'], ['Registriernummer/MRN', 'Registriernummer / MRN', 'MRN', 'Registrienummer/MRN'], required=not data_sources['df_import_zl'].empty) if not data_sources['df_import_zl'].empty else None,
            'pos_field_eza': find_col(data_sources['df_import_eza'], ['PositionNo'], required=not data_sources['df_import_eza'].empty) if not data_sources['df_import_eza'].empty else None,
            'pos_field_zl': find_col(data_sources['df_import_zl'], ['PositionNo'], required=not data_sources['df_import_zl'].empty) if not data_sources['df_import_zl'].empty else None,
            'ncts_mrn_col': 'MRN' if not data_sources['df_ncts'].empty and 'MRN' in data_sources['df_ncts'].columns else None,
            'suma_pos_col': None
        }
        
        suma_pos_candidates = ['Position SumA', 'Pos. SumA', 'PositionNo SumA', 'Position', 'Pos', 'PositionNo']
        for candidate in suma_pos_candidates:
            if candidate in df_leit.columns:
                field_mappings['suma_pos_col'] = candidate
                break
        
        if not field_mappings['suma_pos_col']:
            st.warning("⚠️ SUMA-Position-Spalte nicht gefunden. Verwende leeres Feld.")
        
        update_progress(progress_bar, 5, 100, "Bereite Daten vor")
        
        with schritte_container:
            st.success("✅ Daten vorbereitet und MRN-Werte bereinigt")
        
        df_leit[field_mappings['leit_col_weitere']] = df_leit[field_mappings['leit_col_weitere']].apply(clean_mrn)
        df_leit[field_mappings['leit_col_reg']] = df_leit[field_mappings['leit_col_reg']].apply(clean_mrn)
        
        for source, col in [('df_import_eza', 'import_eza_col'), ('df_import_zl', 'import_zl_col'), ('df_ncts', 'ncts_mrn_col')]:
            if field_mappings[col] and not data_sources[source].empty:
                data_sources[source][field_mappings[col]] = data_sources[source][field_mappings[col]].apply(clean_mrn)
        
        stats = defaultdict(int)
        
        results = []
        
        total_steps = len(VERARBEITBARE_ARTEN) + len(PAUSCHALE_ARTEN)
        
        for i, anmeldeart in enumerate(VERARBEITBARE_ARTEN):
            count = df_leit[field_mappings['anmeldeart_col']].eq(anmeldeart).sum()
            
            if count > 0:
                update_progress(progress_bar, i + 1, total_steps, "Verarbeite", f"{anmeldeart}-Anmeldearten ({count} Zeilen)")
                
                if anmeldeart in ['IMDC', 'WIDS']:
                    import_source = 'df_import_eza' if anmeldeart == 'IMDC' else 'df_import_zl'
                    if not data_sources[import_source].empty:
                        results.extend(process_anmeldeart_generic(
                            anmeldeart, df_leit, data_sources, field_mappings, stats
                        ))
                        with schritte_container:
                            st.success(f"✅ {anmeldeart}-Anmeldearten verarbeitet ({count} Zeilen)")
                elif anmeldeart == 'IPDC':
                    results.extend(process_anmeldeart_generic(
                        anmeldeart, df_leit, data_sources, field_mappings, stats
                    ))
                    with schritte_container:
                        st.success(f"✅ {anmeldeart}-Anmeldearten verarbeitet ({count} Zeilen)")
                elif anmeldeart == 'NCDP' and not data_sources['df_ncts'].empty:
                    results.extend(process_anmeldeart_generic(
                        anmeldeart, df_leit, data_sources, field_mappings, stats
                    ))
                    with schritte_container:
                        st.success(f"✅ {anmeldeart}-Anmeldearten verarbeitet ({count} Zeilen)")
            else:
                update_progress(progress_bar, i + 1, total_steps, "", f"Keine {anmeldeart}-Anmeldearten vorhanden")
        
        pauschale_map = {
            '(leer)': None,
            'APDC': 'APDC',
            'AVDC': 'AVDC',
            'NCAR': 'NCAR'
        }
        
        for j, (anmeldeart_name, anmeldeart_filter) in enumerate(pauschale_map.items()):
            count = st.session_state.stats.get(anmeldeart_name, 0)
            if count > 0:
                update_progress(progress_bar, len(VERARBEITBARE_ARTEN) + j + 1, total_steps, 
                               "Verarbeite", f"{anmeldeart_name}-Anmeldearten ({count} Zeilen)")
                results.extend(process_pauschale_anmeldeart(
                    df_leit, field_mappings, stats, anmeldeart_filter, anmeldeart_name
                ))
                with schritte_container:
                    st.success(f"✅ {anmeldeart_name}-Anmeldearten verarbeitet ({count} Zeilen)")
            else:
                update_progress(progress_bar, len(VERARBEITBARE_ARTEN) + j + 1, total_steps, 
                               "", f"Keine {anmeldeart_name}-Anmeldearten vorhanden")
        
        update_progress(progress_bar, 95, 100, "Wende Geschäftsregeln an")
        
        st.session_state['atb_filtered_count'] = stats.get('atb_skipped', 0)
        
        results = apply_zoelle_rule(results)
        
        with schritte_container:
            st.success("✅ Geschäftsregeln angewendet (Mindestabgaben, Pauschalen)")
        
        update_progress(progress_bar, 100, 100, "✅ Verarbeitung abgeschlossen")
        time.sleep(0.5)
        
        if results:
            ziel = pd.DataFrame(results)
            ziel = prepare_dataframe_for_sorting(ziel)
            ziel_sorted = sort_dataframe_standard(ziel).reset_index(drop=True)
            
            with schritte_container:
                st.success(f"✅ Ergebnis erstellt: {len(ziel_sorted)} Zeilen")

            st.session_state['ziel_sorted'] = ziel_sorted
            st.session_state['processing_stats'] = dict(stats)
            st.session_state['results_available'] = True
            
            st.success(f"✅ Verarbeitung erfolgreich abgeschlossen! {len(ziel_sorted)} Zeilen erstellt.")
           
            display_results(ziel_sorted, dict(stats))
            
            process_buergschaft(ziel_sorted)
        else:
            st.warning("⚠️ Keine Daten zum Verarbeiten gefunden.")
            
    except Exception as e:
        error_msg = f"{type(e).__name__}: {str(e)}"
        st.session_state['processing_error'] = error_msg
        st.error(f"❌ Fehler bei der Verarbeitung: {error_msg}")
        
        with st.expander("🔍 Debug-Informationen", expanded=False):
            st.code(traceback.format_exc())
            
    finally:
        st.session_state['processing_active'] = False
        
        if 'progress_bar' in locals():
            progress_bar.empty()

        if 'results' in locals() and results and not st.session_state.get('processing_error'):
            time.sleep(0.1)

# === ERGEBNIS-ANZEIGE ===

def display_results(ziel, stats):
    """Zeigt Verarbeitungsergebnisse an"""
    st.markdown("---")
    st.subheader("6. Ergebnisse", help="Übersicht aller berechneten Werte für die Bürgschaftsbelastung. Hier sehen Sie die Zusammenfassung aller verarbeiteten Vorgänge.")
    
    st.subheader("6.1 Verarbeitungsergebnisse", help="Anzahl der erstellten Zeilen pro Anmeldeart. 'KEIN MATCH' bedeutet, dass keine passenden Importdaten gefunden wurden und daher der Pauschalbetrag verwendet wurde.")
    st.success(f"✅ Verarbeitung erfolgreich! {len(ziel)} Zeilen erstellt.")
    
    if st.session_state.get('atb_filtered_count', 0) > 0:
        st.info(f"""
        ℹ️ **ATB-Filter:** {st.session_state['atb_filtered_count']} Zeilen mit ATB in 'Weitere Registriernummer Folgeverfahren' 
        wurden übersprungen (S-Anmeldearten und andere interne Vorgänge).
        """)
    
    s_arten_summe = sum(st.session_state.stats.get(art, 0) for art in S_ANMELDEARTEN)
    if s_arten_summe > 0:
        st.info(f"""
        ℹ️ **Hinweis:** {s_arten_summe} S-Anmeldearten (SUSP, SUDC, SUCO, SUCF) wurden in der Leitdatei gefunden, 
        aber nicht verarbeitet, da sie keine Bürgschaftsrelevanz haben (interne Konsolidierungen/Aufteilungen).
        """)
    
    col1, col2, col3, col4 = st.columns(4)
    col5, col6, col7, col8 = st.columns(4)
    
    metrics = [
        ("Gesamtzeilen", len(ziel)),
        ("IMDC-Zeilen", len(ziel[ziel['Anmeldeart'] == 'IMDC'])),
        ("WIDS-Zeilen", len(ziel[ziel['Anmeldeart'] == 'WIDS'])),
        ("IPDC-Zeilen", len(ziel[ziel['Anmeldeart'] == 'IPDC'])),
        ("NCDP-Zeilen", len(ziel[ziel['Anmeldeart'] == 'NCDP'])),
        ("(leer)-Zeilen", len(ziel[ziel['Anmeldeart'] == '(leer)'])),
        ("APDC-Zeilen", len(ziel[ziel['Anmeldeart'] == 'APDC'])),
        ("AVDC-Zeilen", len(ziel[ziel['Anmeldeart'] == 'AVDC']))
    ]
    
    for col, (label, value) in zip([col1, col2, col3, col4, col5, col6, col7, col8], metrics):
        with col:
            st.metric(label, value)
    
    ncar_count = len(ziel[ziel['Anmeldeart'] == 'NCAR'])
    if ncar_count > 0:
        st.metric("NCAR-Zeilen", ncar_count)
    
    display_processing_protocol(stats)
    display_financial_summary(ziel)

def display_processing_protocol(stats):
    """Zeigt Verarbeitungsprotokoll als kompakte Tabelle"""
    st.subheader("6.2 Verarbeitungsprotokoll", help="Detaillierte Aufschlüsselung der Verarbeitung: Wie viele Positionen wurden in den Kalkulationsdateien gefunden (Mit Match) und wie viele mussten mit dem Pauschalbetrag berechnet werden (Ohne Match).")
    with st.expander("📊 Details anzeigen", expanded=True):
        protocol_data = []
        
        if st.session_state.stats.get("IMDC", 0) > 0:
            protocol_data.append({
                'Anmeldeart': 'IMDC',
                'Leitdatei': st.session_state.stats.get("IMDC", 0),
                'Verarbeitet': stats.get('processed_imdc', 0),
                'Mit Match': stats.get('imdc_match', 0),
                'Ohne Match': stats.get('imdc_no_match', 0),
                'Details': f"{stats.get('imdc_3criteria_match', 0)} präzise, {stats.get('imdc_fallback_match', 0)} Fallback"
            })
        
        if st.session_state.stats.get("WIDS", 0) > 0:
            protocol_data.append({
                'Anmeldeart': 'WIDS',
                'Leitdatei': st.session_state.stats.get("WIDS", 0),
                'Verarbeitet': stats.get('processed_wids', 0),
                'Mit Match': stats.get('wids_match', 0),
                'Ohne Match': stats.get('wids_no_match', 0),
                'Details': st.session_state.get('wids_aggregation', 'Position mit höchstem Zollwert')
            })
        
        if st.session_state.stats.get("IPDC", 0) > 0:
            protocol_data.append({
                'Anmeldeart': 'IPDC',
                'Leitdatei': st.session_state.stats.get("IPDC", 0),
                'Verarbeitet': stats.get('processed_ipdc', 0),
                'Mit Match': '—',
                'Ohne Match': '—',
                'Details': f"{stats.get('ipdc_with_zollwert', 0)} mit Zollwert, {stats.get('ipdc_without_zollwert', 0)} ohne"
            })
        
        if st.session_state.stats.get("NCDP", 0) > 0:
            protocol_data.append({
                'Anmeldeart': 'NCDP',
                'Leitdatei': st.session_state.stats.get("NCDP", 0),
                'Verarbeitet': stats.get('processed_ncdp', 0),
                'Mit Match': stats.get('ncdp_match', 0),
                'Ohne Match': stats.get('ncdp_no_match', 0),
                'Details': 'Sicherheitsbetrag aus NCTS'
            })
        
        if st.session_state.stats.get("(leer)", 0) > 0:
            protocol_data.append({
                'Anmeldeart': '(leer)',
                'Leitdatei': st.session_state.stats.get("(leer)", 0),
                'Verarbeitet': stats.get('(leer)_processed', 0),
                'Mit Match': '—',
                'Ohne Match': '—',
                'Details': st.session_state.get('leere_anmeldeart_option', 'Alle Zeilen (wie andere Anmeldearten)')
            })
        
        for art in ['APDC', 'AVDC']:
            if st.session_state.stats.get(art, 0) > 0:
                protocol_data.append({
                    'Anmeldeart': art,
                    'Leitdatei': st.session_state.stats.get(art, 0),
                    'Verarbeitet': stats.get(f'{art.lower()}_processed', 0),
                    'Mit Match': '—',
                    'Ohne Match': '—',
                    'Details': 'Pauschale (wie leere)'
                })
        
        if st.session_state.stats.get("NCAR", 0) > 0:
            protocol_data.append({
                'Anmeldeart': 'NCAR',
                'Leitdatei': st.session_state.stats.get("NCAR", 0),
                'Verarbeitet': stats.get('ncar_processed', 0),
                'Mit Match': '—',
                'Ohne Match': '—',
                'Details': f'Pauschale {st.session_state.get("pauschalbetrag", 10000):.0f}€'
            })
        
        if protocol_data:
            df_protocol = pd.DataFrame(protocol_data)
            st.dataframe(df_protocol, hide_index=True, use_container_width=True)
            
            if stats.get('imdc_be_anteil_rows', 0) > 0:
                be_prozent = (stats['imdc_be_anteil_rows'] / stats.get('processed_imdc', 1) * 100)
                st.info(f"💡 BE-Anteil: {stats['imdc_be_anteil_rows']} Zeilen ({be_prozent:.1f}%) mit BE-Anteil-Info verarbeitet")
        
        if stats.get('atb_skipped', 0) > 0:
            st.caption(f"🔵 {stats['atb_skipped']} Zeilen mit ATB in 'Weitere Registriernummer' wurden übersprungen")
        
        s_arten_summe = sum(st.session_state.stats.get(art, 0) for art in S_ANMELDEARTEN)
        if s_arten_summe > 0:
            st.caption(f"⚫ {s_arten_summe} S-Anmeldearten (interne Konsolidierungen) wurden nicht verarbeitet")

def display_financial_summary(ziel):
    """Zeigt finanzielle Zusammenfassung"""
    st.subheader("6.3 Finanzielle Zusammenfassung", help="Gesamtsummen aller Zollwerte und Abgaben. Diese Werte bilden die Grundlage für die Bürgschaftsbelastung. Die EUSt wird separat ausgewiesen und ist nicht Teil der Bürgschaftsberechnung.")
    
    s_arten_summe = sum(st.session_state.stats.get(art, 0) for art in S_ANMELDEARTEN)
    if s_arten_summe > 0:
        st.warning(f"""
        ⚠️ **Hinweis:** {s_arten_summe} S-Anmeldearten wurden nicht in die finanzielle Zusammenfassung einbezogen, 
        da sie keine Bürgschaftsrelevanz haben (interne Vorgänge).
        """)
    
    ziel_numeric = ziel.copy()
    ziel_numeric['Zollwert (total)'] = pd.to_numeric(ziel_numeric['Zollwert (total)'], errors='coerce').fillna(0)
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Gesamt-Zollwert", f"€ {ziel_numeric['Zollwert (total)'].sum():,.2f}")
    with col2:
        st.metric("Gesamt-Zölle", f"€ {ziel['Zölle (total)'].sum():,.2f}")
    with col3:
        st.metric("Gesamt-EUSt", f"€ {ziel['EUSt'].sum():,.2f}")
    with col4:
        st.metric("Gesamtabgaben", f"€ {ziel['Gesamtabgaben'].sum():,.2f}")
    
    st.subheader("📊 Aufschlüsselung nach Anmeldeart")
    
    anmeldearten = VERARBEITBARE_ARTEN + PAUSCHALE_ARTEN
    
    summary_data = []
    
    for art in anmeldearten:
        art_data = ziel[ziel['Anmeldeart'] == art]
        if len(art_data) > 0:
            art_numeric = art_data.copy()
            art_numeric['Zollwert (total)'] = pd.to_numeric(art_numeric['Zollwert (total)'], errors='coerce').fillna(0)
            
            summary_data.append({
                'Anmeldeart': art,
                'Anzahl': len(art_data),
                'Zollwert': art_numeric['Zollwert (total)'].sum(),
                'Zölle': art_data['Zölle (total)'].sum(),
                'EUSt': art_data['EUSt'].sum(),
                'Gesamtabgaben': art_data['Gesamtabgaben'].sum()
            })
    
    if summary_data:
        summary_df = pd.DataFrame(summary_data)
        
        formatted_df = summary_df.copy()
        for col in ['Zollwert', 'Zölle', 'EUSt', 'Gesamtabgaben']:
            formatted_df[col] = formatted_df[col].apply(lambda x: f"€ {x:,.2f}")
        
        st.dataframe(formatted_df, hide_index=True, use_container_width=True)
        
        st.markdown("---")
        col1, col2, col3, col4, col5, col6 = st.columns(6)
        with col1:
            st.write("**GESAMT**")
        with col2:
            st.write(f"**{summary_df['Anzahl'].sum()}**")
        with col3:
            st.write(f"**€ {summary_df['Zollwert'].sum():,.2f}**")
        with col4:
            st.write(f"**€ {summary_df['Zölle'].sum():,.2f}**")
        with col5:
            st.write(f"**€ {summary_df['EUSt'].sum():,.2f}**")
        with col6:
            st.write(f"**€ {summary_df['Gesamtabgaben'].sum():,.2f}**")
    
    st.info(f"""
    ℹ️ **Hinweise:** - Gesamtabgaben zwischen 0,01€ und 0,99€ werden auf 1€ angehoben
    - Bei Gesamtabgaben = 0€: Wenn Zollwert > 0 → 1€, sonst {st.session_state.get('pauschalbetrag', 10000):.0f}€ Pauschale
    - Zölle sind IMMER gleich Gesamtabgaben
    - Die EUSt wird separat ausgewiesen und ist NICHT in den Gesamtabgaben enthalten
    """)

def process_buergschaft(ziel):
    """Verarbeitet Bürgschaftssaldo-Berechnung"""
    st.subheader("6.4 Bürgschaftssaldo-Berechnung", help="Chronologische Darstellung aller Ein- und Ausgänge mit täglichen Salden. Zeigt die Entwicklung der Bürgschaftsauslastung über den gesamten Zeitraum mit Höchst- und Tiefstständen.")
    
    with st.spinner("💰 Bürgschaftssaldo wird berechnet..."):
        time.sleep(0.3)
        bewegungen_df = create_bewegungstabelle(ziel)
        daily_summary = calculate_daily_summary(bewegungen_df, st.session_state.get('startbuergschaft', 0))
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Startbürgschaft", f"€ {st.session_state.get('startbuergschaft', 0):,.2f}")
        
        total_belastung = sum(d['Belastung'] for d in daily_summary.values())
        total_entlastung = sum(d['Entlastung'] for d in daily_summary.values())
        end_stand = st.session_state.get('startbuergschaft', 0) - total_belastung + total_entlastung
        
        with col2:
            st.metric("Gesamtbelastung", f"€ {total_belastung:,.2f}")
            st.metric("Gesamtentlastung", f"€ {total_entlastung:,.2f}")
        
        with col3:
            st.metric("Endbürgschaft", f"€ {end_stand:,.2f}")
            auslastung = 0 if st.session_state.get('startbuergschaft', 0) == 0 else ((st.session_state.get('startbuergschaft', 0) - end_stand) / st.session_state.get('startbuergschaft', 0) * 100)
            st.metric("Auslastung", f"{auslastung:.1f}%")
        
        ziel_mit_saldo = add_tagessummen_to_ziel(ziel, daily_summary)
        bewegungsdetails_df = create_bewegungsdetails_df(bewegungen_df, daily_summary, st.session_state.get('startbuergschaft', 0))
        tageszusammenfassung_df = create_tageszusammenfassung_df_mit_extrema(bewegungen_df, daily_summary, st.session_state.get('startbuergschaft', 0))
        
        if len(tageszusammenfassung_df) > 1:
            gesamt_row = tageszusammenfassung_df[tageszusammenfassung_df['Datum'] == 'GESAMT']
            if not gesamt_row.empty:
                max_auslastung = gesamt_row['Auslastung %'].iloc[0]
                tiefststand = gesamt_row['Tiefststand'].iloc[0]
        
            st.session_state['max_auslastung_str'] = f"{max_auslastung:.2f} %".replace('.', ',')
            st.session_state['tiefststand_str'] = format_currency(tiefststand)
            st.session_state['max_auslastung'] = f"{max_auslastung:.1f}%"
            st.session_state['ziel_sorted'] = ziel

        ncar_info = ""
        if st.session_state.get('ncar_enabled', True) and 'df_ncar' in st.session_state and st.session_state['df_ncar'] is not None:
            with st.spinner("✨ NCAR-Daten werden eingearbeitet..."):
                ziel_mit_saldo = enhance_ziel_with_ncar(ziel_mit_saldo, st.session_state['df_ncar'])
                
                transport_mrn_rows = (ziel_mit_saldo['MRN-Nummer Eingang'] != ziel_mit_saldo['ATB-Nummer']).sum()
                packstuck_rows = (pd.to_numeric(ziel_mit_saldo['Menge'], errors='coerce') > 0).sum()
                
                ncar_info = f" (inkl. NCAR: {transport_mrn_rows} Transport-MRN, {packstuck_rows} mit Packstücken)"
        
        buergschaft_info = ""
        if st.session_state.get('buergschaft_erhöhung_aktiv', False):
            betrag = st.session_state.get('buergschaft_erhöhung_betrag', 0)
            datum = st.session_state.get('buergschaft_erhöhung_datum', date.today())
            buergschaft_info = f" | Bürgschaft +{betrag/1000000:.1f} Mio am {datum.strftime('%d.%m.%Y')}"
        
        st.success(f"✅ Bürgschaftssaldo wurde berechnet!{ncar_info}{buergschaft_info}")
        
        st.info(f"""
        **Excel wird 3 Sheets enthalten:**
        1. **Ergebnis** - {len(ziel_mit_saldo)} Zeilen mit Tagessalden{ncar_info}
        2. **Bewegungsdetails** - {len(bewegungsdetails_df)} Zeilen mit allen Ein-/Ausgängen
        3. **Tageszusammenfassung** - {len(tageszusammenfassung_df)} Zeilen mit Höchst-/Tiefstständen pro Tag
        """)
        
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            ziel_mit_saldo.to_excel(writer, index=False, sheet_name='Ergebnis')
            bewegungsdetails_df.to_excel(writer, index=False, sheet_name='Bewegungsdetails')
            tageszusammenfassung_df.to_excel(writer, index=False, sheet_name='Tageszusammenfassung')
        output.seek(0)

        st.session_state['excel_file'] = output.getvalue()
        
        # HISTORIE-HOOK: Nach erfolgreicher Verarbeitung speichern
        save_to_history()

        st.markdown("---")
        st.success("✅ Excel-Datei wurde erfolgreich erstellt! Wechseln Sie zum Downloads-Tab.")

# === UI KOMPONENTEN ===

def setup_sidebar():
    """Konfiguriert die Sidebar mit neuer, aufgeräumter Struktur"""
    
    st.markdown("""
    <style>
        [data-testid="stSidebarNav"] {
            display: none;
        }
    </style>
    """, unsafe_allow_html=True)
    
    with st.sidebar:
        st.markdown("""
        <div style="background: white; padding: 1rem 2rem; border-radius: 12px; border: 3px solid #8B1C1C;">
        """ + render_logo(size="small", with_tagline=True) + """
        </div>
        """, unsafe_allow_html=True)
        
        if st.session_state.get('mandant_logo') and os.path.exists(st.session_state.get('mandant_logo', '')):
            try:
                st.markdown("<br>", unsafe_allow_html=True)  # <-- Abstand hinzufügen
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    st.image(st.session_state['mandant_logo'], width=150)
            except Exception:
                pass
        
        st.markdown(f"### 🏢 {st.session_state.get('mandant', 'Unbekannt')}")
        
        st.markdown("---")
        
        st.subheader("⚙️ Konfiguration", help="Ihre aktuellen Mandanten-Einstellungen für die Bürgschaftsberechnung. Diese Werte bestimmen den Verarbeitungszeitraum und die Berechnungsparameter.")
        
        settings = load_settings()
        if settings and 'current_config' in settings:
            config = settings[settings['current_config']]
            
            st.markdown(f"""
            • **Zeitraum:** {config['von'][:5]} - {config['bis'][3:]}  
            • **Bürgschaft:** {format_currency(config['buergschaft'])}  
            • **Pauschale:** {format_currency(config['pauschale'])}  
            • **Ersatz:** {config['ersatz_zollsatz']}%
            """)
            
            if config.get('buergschaft_erhoehung_aktiv', False):
                betrag = config.get('buergschaft_erhoehung_betrag', 0)
                datum = config.get('buergschaft_erhoehung_datum', 'N/A')
                st.markdown(f"• **Erhöhung:** +{betrag/1000000:.1f} Mio am {datum[:5]}")
        
        st.info("⚙️ Einstellungen finden Sie im Tab **'Einstellungen'** oben")
        
        st.markdown("---")
        
        st.subheader("💡 Geschäftsregeln", help="Automatische Regeln die bei der Verarbeitung angewendet werden, um realistische Bürgschaftswerte zu gewährleisten.")
        st.info("""
        • **Mindestabgabe:** 1 €
        • **Bei 0 €:** Pauschale
        • **Zölle = Gesamtabgaben**
        """)
        
        if st.session_state.get('df_leit') is not None and st.session_state.get('stats'):
            st.markdown("---")
            st.subheader("📊 Geladene Daten", help="Übersicht der geladenen Leitdatei und Verteilung der Anmeldearten. Zeigt wie viele Zeilen pro Anmeldeart verarbeitet werden.")
            
            total = st.session_state.stats.get("Gesamt", 0)
            st.metric("Leitdatei", format_currency(total, display_only=True) + " Zeilen")
            
            anmeldearten = VERARBEITBARE_ARTEN + ['(leer)']
            for art in anmeldearten:
                count = st.session_state.stats.get(art, 0)
                if count > 0:
                    st.caption(f"{art}: {count}")
        
        if st.session_state.get('processing_active', False):
            st.markdown("---")
            st.warning("⏳ Verarbeitung läuft...")

        if 'excel_file' in st.session_state and st.session_state['excel_file'] is not None:
            st.markdown("---")
            st.success("📥 **Downloads bereit!**")
            st.caption("Klicken Sie auf '2️⃣ Downloads' oben")
            
        # HISTORIE-HOOK: Historie-Anzeige in Sidebar
        history = load_history()
        if history:
            st.markdown("---")
            st.info(f"📊 **Bürgschaftsverlauf:** {len(history)} Einträge")

        st.markdown("---")
        
        st.markdown("""
        <div style="background-color: #fee2e2; border: 2px solid #ef4444; border-radius: 8px; padding: 10px; margin-bottom: 10px;">
            <p style="color: #dc2626; margin: 0; font-weight: bold; text-align: center;">
                ⚠️ ACHTUNG: Alle hochgeladenen Daten werden gelöscht!
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        if 'confirm_reset' not in st.session_state:
            st.session_state.confirm_reset = False
        
        if not st.session_state.confirm_reset:
            if st.button("❌ Neu starten", type="secondary", use_container_width=True,
                        help="Löscht alle Daten und startet neu"):
                st.session_state.confirm_reset = True
                st.rerun()
        else:
            st.warning("Wirklich alle Daten löschen?")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("✅ Ja, alles löschen", type="primary", use_container_width=True):
                    keys_to_keep = ['authenticated', 'mandant', 'mandant_logo', 'mandant_data']
                    for key in list(st.session_state.keys()):
                        if key not in keys_to_keep:
                            del st.session_state[key]
                    st.rerun()
            with col2:
                if st.button("❌ Abbrechen", type="secondary", use_container_width=True):
                    st.session_state.confirm_reset = False
                    st.rerun()        
    
        st.markdown("---")
        st.caption("© Scills GmbH 2025")

def show_file_upload_section():
    """Zeigt die Datei-Upload-Sektion"""
    if 'mandant' in st.session_state:
        st.caption(f"Mandant: {st.session_state['mandant']}")
    
    st.subheader("1. Leitdatei (SumA) hochladen", help="Excel-Export aus SumA mit allen Verwahrungsvorgängen. Diese Datei enthält alle Bewegungen und bildet die Grundlage für die Bürgschaftsberechnung.")
    
    if 'df_leit_unfiltered' in st.session_state and st.session_state.df_leit_unfiltered is not None:
        additional_info = []
        if 'df_leit' in st.session_state and st.session_state.df_leit is not None:
            additional_info.append(f"Gefiltert: {len(st.session_state.df_leit)} Zeilen im Bewilligungszeitraum")
        
        if show_file_status(f"Leitdatei geladen: {len(st.session_state.df_leit_unfiltered)} Zeilen", 
                          len(st.session_state.df_leit_unfiltered), 
                          "reload_leitdatei",
                          None,
                          additional_info):
            keys_to_delete = ['df_leit', 'df_leit_unfiltered', 'leitdatei_bytes', 
                            'stats', 'datum_filter_confirmed', 'df_ncar', 'ncar_bytes',
                            'df_import_eza', 'df_import_eza_bytes',
                            'df_import_zl', 'df_import_zl_bytes',
                            'df_ncts', 'df_ncts_bytes']
            for key in keys_to_delete:
                if key in st.session_state:
                    del st.session_state[key]
            st.rerun()
        
        show_date_filter_and_imports()
    else:
        leitdatei = st.file_uploader("Leitdatei", type=["xlsx", "xls"], key="leitdatei",
                                    help="Wählen Sie die SumA-Leitdatei aus. Format: Excel (.xlsx oder .xls)")
        
        if leitdatei is not None:
            process_leitdatei(leitdatei)

    if st.session_state.get('results_available', False) and 'ziel_sorted' in st.session_state:
        st.markdown("---")
        
        with st.expander("📊 Ergebnisse erneut anzeigen", expanded=False):
            display_results(st.session_state['ziel_sorted'], st.session_state.get('processing_stats', {}))
        
        if 'excel_file' in st.session_state:
            st.info("✅ Excel-Datei bereits erstellt. Wechseln Sie zum **Downloads-Tab** um sie herunterzuladen.")

def show_date_filter_and_imports():
    """Zeigt Datumsfilter und Import-Uploads"""
    st.subheader("2. Bürgschaftszeitraum auswählen", help="Filtern Sie die Leitdatei auf den bewilligten Zeitraum. Nur Vorgänge mit Gestellungsdatum in diesem Zeitraum werden verarbeitet.")
    
    col1, col2, col3 = st.columns([2, 2, 1])
    with col1:
        von_datum = st.date_input("Von-Datum", value=st.session_state.get('von_datum', date(2023, 11, 1)),
                                 help="Startdatum des Bewilligungszeitraums")
    with col2:
        bis_datum = st.date_input("Bis-Datum", value=st.session_state.get('bis_datum', date(2024, 4, 30)),
                                 help="Enddatum des Bewilligungszeitraums")
    
    st.session_state['von_datum'] = von_datum
    st.session_state['bis_datum'] = bis_datum
    
    with col3:
        st.markdown("<br>", unsafe_allow_html=True)
        filter_button = st.button("🔍 Filter anwenden", use_container_width=True, help="Filtert die Leitdatei auf den gewählten Zeitraum")
    
    if filter_button or st.session_state.get('datum_filter_confirmed', False):
        st.session_state['datum_filter_confirmed'] = True
        
        df_leit = st.session_state.df_leit_unfiltered
        gestell_col = find_col(df_leit, ['Datum Überlassung - CUSTST'])
        df_leit['_gestell_date'] = pd.to_datetime(df_leit[gestell_col], errors='coerce').dt.date
        df_leit = df_leit.dropna(subset=['_gestell_date'])
        mask = (df_leit['_gestell_date'] >= von_datum) & (df_leit['_gestell_date'] <= bis_datum)
        df_leit_filtered = df_leit[mask].copy()
        df_leit_filtered = df_leit_filtered.drop(columns=['_gestell_date'])
        
        if df_leit_filtered.empty:
            st.warning("⚠️ Keine Daten im gewählten Zeitraum gefunden!")
            return
        
        st.session_state.df_leit = df_leit_filtered
        
        anmeldeart_col = find_col(df_leit_filtered, ['Anmeldeart Folgeverfahren'])
        st.session_state.stats = calculate_statistics(df_leit_filtered, anmeldeart_col)
        
        st.success(f"✅ {len(df_leit_filtered)} Einträge im Zeitraum {von_datum.strftime('%d.%m.%Y')} - {bis_datum.strftime('%d.%m.%Y')}")
        display_statistics_table()
        
        show_ncar_upload()
        
        show_all_imports()
        show_processing_button()
    else:
        st.info("👆 Bitte wählen Sie einen Zeitraum und klicken Sie auf 'Filter anwenden'")           

def process_leitdatei(leitdatei):
    """Verarbeitet die Leitdatei - NUR wenn noch nicht geladen"""
    if 'df_leit' in st.session_state and st.session_state.df_leit is not None:
        return
    
    try:
        file_bytes = leitdatei.getvalue()
        st.session_state['leitdatei_bytes'] = file_bytes
        
        with st.spinner("Leitdatei wird geladen..."):
            df_leit = pd.read_excel(io.BytesIO(file_bytes))
            
            required_leit_cols = [
                ['Datum Überlassung - CUSTST'],
                ['Weitere Registriernummer Folgeverfahren', 'Weitere Registriernummer'],
                ['Registriernummer Folgeverfahren'],
                ['Anmeldeart Folgeverfahren'],
                ['Bezugsnummer/LRN SumA'],
                ['Registriernummer/MRN SumA'],
                ['Datum Ende - CUSFIN']
            ]
            
            validate_dataframe(df_leit, required_leit_cols, "Leitdatei")
        
        st.session_state.df_leit_unfiltered = df_leit
        st.success(f"✅ Leitdatei erfolgreich hochgeladen ({len(df_leit)} Zeilen)")
        st.rerun()
        
    except Exception as e:
        st.error(f"❌ Fehler beim Laden der Leitdatei: {e}")
        if 'leitdatei_bytes' in st.session_state:
            del st.session_state['leitdatei_bytes']

def show_ncar_upload():
    """Zeigt NCAR-Upload mit verstecktem Uploader nach erfolgreichem Upload"""
    if not st.session_state.get('ncar_enabled', True):
        return
    
    st.markdown("---")
    st.subheader("3. NCAR-Versandbeendigung (NCTS TBE)", help="Optional: Laden Sie NCAR-Daten hoch, um Transport-MRN und Packstückzahlen automatisch in der Ergebnisliste zu ergänzen. Verbessert die Datenqualität für Zollprüfungen.")
    
    if 'df_ncar' in st.session_state and st.session_state['df_ncar'] is not None:
        additional_info = ["Transport-MRN und Packstücke werden automatisch ergänzt"]
        
        if show_file_status(f"NCAR-Datei geladen", 
                          len(st.session_state.df_ncar), 
                          "reload_ncar",
                          "df_ncar",
                          additional_info):
            st.rerun()
    else:
        st.info("Laden Sie eine NCAR-Datei hoch, um Transport-MRN und Packstücke automatisch zu ergänzen")
        
        ncar_file = st.file_uploader(
            "NCAR-Datei", 
            type=["xlsx", "xls"], 
            key="ncar_upfront",
            help="NCTS-Export mit Versandbeendigungen. Benötigte Spalten: Registriernr.-SumA, RegistriernNr./MRN, Anzahl Packstücke"
        )
        
        if ncar_file:
            file_bytes = ncar_file.getvalue()
            st.session_state['ncar_bytes'] = file_bytes
            
            with st.spinner("NCAR-Datei wird verarbeitet..."):
                try:
                    ncar_df = pd.read_excel(io.BytesIO(file_bytes))
                    
                    required_cols = ['Registriernr.-SumA', 'RegistriernNr./MRN', 'Anzahl Packstücke']
                    
                    if not validate_import_file(ncar_df, required_cols, "NCAR-Datei"):
                        if 'ncar_bytes' in st.session_state:
                            del st.session_state['ncar_bytes']
                    else:
                        st.session_state['df_ncar'] = ncar_df
                        st.rerun()
                        
                except Exception as e:
                    st.error(f"❌ Fehler beim Lesen der NCAR-Datei: {e}")
                    if 'ncar_bytes' in st.session_state:
                        del st.session_state['ncar_bytes']

def display_statistics_table():
    """Zeigt Statistik-Tabelle"""
    verarbeitbare_arten = VERARBEITBARE_ARTEN + ['APDC', 'AVDC', 'NCAR']
    
    verarbeitbar_data = []
    intern_data = []
    
    for art in verarbeitbare_arten:
        if st.session_state.stats.get(art, 0) > 0:
            verarbeitbar_data.append({
                "Anmeldeart": art, 
                "Zeilen": st.session_state.stats.get(art, 0),
                "Status": "✅"
            })
    
    for art in S_ANMELDEARTEN:
        if st.session_state.stats.get(art, 0) > 0:
            intern_data.append({
                "Anmeldeart": art, 
                "Zeilen": st.session_state.stats.get(art, 0),
                "Status": "⚫"
            })
    
    if st.session_state.stats.get("(leer)", 0) > 0:
        verarbeitbar_data.append({
            "Anmeldeart": "(leer)", 
            "Zeilen": st.session_state.stats.get("(leer)", 0),
            "Status": "✅"
        })
    
    all_data = verarbeitbar_data + intern_data
    
    if all_data:
        col1, col2 = st.columns([3, 1])
        with col1:
            stats_df = pd.DataFrame(all_data)
            st.dataframe(stats_df, hide_index=True, use_container_width=True)
        
        with col2:
            st.metric("Gesamt", st.session_state.stats.get("Gesamt", 0))
            
            verarbeitbar_summe = sum(item["Zeilen"] for item in verarbeitbar_data)
            intern_summe = sum(item["Zeilen"] for item in intern_data)
            
            st.write("**Aufschlüsselung:**")
            st.write(f"- Verarbeitbar: **{verarbeitbar_summe}**")
            if intern_summe > 0:
                st.write(f"- Intern (S-Arten): **{intern_summe}**")
    
    for anmeldeart, file_type in [("IMDC", "Importdatei EZA"), ("WIDS", "Importdatei ZL"), ("NCDP", "NCTS-Datei")]:
        if st.session_state.stats.get(anmeldeart, 0) > 0:
            anmeldeart_col = find_col(st.session_state.df_leit, ['Anmeldeart Folgeverfahren'])
            anmeldeart_data = st.session_state.df_leit[st.session_state.df_leit[anmeldeart_col] == anmeldeart]
            
            erledigungs_daten = pd.to_datetime(anmeldeart_data['Datum Ende - CUSFIN'], errors='coerce').dropna()
            
            if not erledigungs_daten.empty:
                von_datum = erledigungs_daten.min()
                bis_datum = erledigungs_daten.max()
                
                von_str = von_datum.strftime('%m/%Y')
                bis_str = bis_datum.strftime('%m/%Y')
                
                if von_str == bis_str:
                    zeitraum_str = von_str
                else:
                    zeitraum_str = f"{von_str} - {bis_str}"
                
                st.info(f"📋 Für {st.session_state.stats.get(anmeldeart, 0)} {anmeldeart}-Zeilen wird eine {file_type} benötigt **({zeitraum_str})**")
            else:
                st.info(f"📋 Für {st.session_state.stats.get(anmeldeart, 0)} {anmeldeart}-Zeilen wird eine {file_type} benötigt.")

def show_all_imports():
    """Zeigt alle Import-Uploads nebeneinander"""
    st.markdown("---")
    st.subheader("4. Kalkulationsdateien hochladen", help="Laden Sie die Kalkulationsdateien mit Zollwerten und Abgaben hoch. Nur benötigte Dateien werden aktiviert - basierend auf den gefundenen Anmeldearten in der Leitdatei.")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        process_import_upload(
            anmeldeart="IMDC",
            file_type="4.1 Import EZA",
            file_key="importdatei_eza",
            session_key="df_import_eza",
            required_cols=[
                ['Registriernummer/MRN', 'Registriernummer / MRN', 'MRN'],
                ['PositionNo'],
                ['Warentarifnummer'],
                ['Zollwert'],
                ['AbgabeZollsatz']
            ],
            special_processing="eza"
        )
    
    with col2:
        process_import_upload(
            anmeldeart="WIDS",
            file_type="4.2 Import ZL / VAV",
            file_key="importdatei_zl",
            session_key="df_import_zl",
            required_cols=[
                ['Registriernummer/MRN', 'Registriernummer / MRN', 'MRN', 'Registrienummer/MRN'],
                ['PositionNo'],
                ['Warentarifnummer'],
                ['Vorraussichtliche Zollabgabe', 'Voraussichtliche Zollabgabe'],
                ['Vorraussichtliche Zollsatzabgabe', 'Voraussichtliche Zollsatzabgabe'],
                ['DV1UmgerechnerterRechnungsbetrag']
            ],
            special_processing=None
        )
    
    with col3:
        process_import_upload(
            anmeldeart="NCDP",
            file_type="4.3 NCTS",
            file_key="nctsdatei",
            session_key="df_ncts",
            required_cols=[
                ['MRN'],
                ['Sicherheit']
            ],
            special_processing=None
        )

def process_import_upload(anmeldeart, file_type, file_key, session_key, required_cols, special_processing=None):
    """Generische Funktion für Import-Uploads mit verstecktem Uploader nach erfolgreicher Verarbeitung"""
    count = st.session_state.stats.get(anmeldeart, 0)
    
    if count > 0:
        help_texts = {
            "IMDC": "Einfuhrzollanmeldungen. Enthält Zollwerte, Zollsätze und Warennummern für IMDC-Vorgänge.",
            "WIDS": "Zolllager-/VAV-Export mit voraussichtlichen Abgaben. Format: Excel aus ATLAS.",
            "NCDP": "NCTS-Export mit Sicherheitsbeträgen für Transitvorgänge. Benötigt MRN und Sicherheit-Spalte."
        }
        
        if session_key in st.session_state and st.session_state[session_key] is not None:
            with st.container():
                st.write(f"**{file_type}** ({count} {anmeldeart})")
                
                df = st.session_state[session_key]
                additional_info = []
                
                if special_processing == "eza" and hasattr(df, 'attrs'):
                    if 'removed_duplicates' in df.attrs:
                        additional_info.append(f"Duplikate entfernt: {df.attrs['removed_duplicates']}")
                    if 'be_multiplied' in df.attrs:
                        additional_info.append(f"BE-Anteil verarbeitet: +{df.attrs['be_multiplied']} Zeilen")
                
                if show_file_status("Erfolgreich geladen", 
                                  len(df), 
                                  f"reload_{file_key}",
                                  session_key,
                                  additional_info):
                    st.rerun()
        
        else:
            st.write(f"**{file_type}** ({count} {anmeldeart})", help=help_texts.get(anmeldeart, ""))
            
            uploaded_file = st.file_uploader(
                file_type.split()[-1],
                type=["xlsx", "xls"],
                key=file_key,
                help="Excel-Datei auswählen (.xlsx oder .xls)"
            )
            
            if uploaded_file is not None:
                file_bytes = uploaded_file.getvalue()
                st.session_state[f"{session_key}_bytes"] = file_bytes
                
                with st.spinner(f"🔄 {file_type} wird verarbeitet..."):
                    try:
                        df_import = pd.read_excel(io.BytesIO(file_bytes))
                        
                        if special_processing == "eza":
                            original_col_count = len(df_import.columns)
                            original_row_count = len(df_import)
                            
                            if st.session_state.get('eza_auto_reduce', True) and original_col_count > len(EXAKTE_EZA_SPALTEN):
                                found_columns = [col for col in EXAKTE_EZA_SPALTEN if col in df_import.columns]
                                if len(found_columns) >= 5:
                                    df_import = df_import[found_columns].copy()
                            
                            if len(df_import.columns) >= 6:
                                col_E = df_import.columns[4]
                                col_F = df_import.columns[5]
                                df_import = df_import.drop_duplicates(subset=[col_E, col_F], keep='first')
                                removed_count = original_row_count - len(df_import)
                            else:
                                removed_count = 0
                            
                            unique_count = len(df_import)
                            
                            df_import = process_eza_be_anteil(df_import)
                            multiplied_count = len(df_import) - unique_count
                            
                            df_import.attrs['removed_duplicates'] = removed_count
                            df_import.attrs['be_multiplied'] = multiplied_count
                        
                        validate_dataframe(df_import, required_cols, file_type)
                        
                        st.session_state[session_key] = df_import
                        
                        st.rerun()
                        
                    except Exception as e:
                        st.error(f"❌ Fehler in {file_type}: {e}")
                        if f"{session_key}_bytes" in st.session_state:
                            del st.session_state[f"{session_key}_bytes"]
    
    else:
        st.write(f"**{file_type}**", help="Diese Datei wird nicht benötigt, da keine entsprechenden Anmeldearten in der Leitdatei gefunden wurden.")
        st.info("Nicht benötigt")
        st.session_state[session_key] = None

def show_processing_button():
    """Zeigt Verarbeitungs-Button wenn alle Dateien vorhanden"""
    if (st.session_state.df_import_eza is not None or st.session_state.df_import_zl is not None or 
        (st.session_state.stats.get("IMDC", 0) == 0 and st.session_state.stats.get("WIDS", 0) == 0) or 
        st.session_state.stats.get("IPDC", 0) > 0 or 
        st.session_state.stats.get("(leer)", 0) > 0 or
        st.session_state.stats.get("APDC", 0) > 0 or
        st.session_state.stats.get("AVDC", 0) > 0 or
        st.session_state.stats.get("NCAR", 0) > 0):
        
        st.markdown("---")
        st.subheader("5. Datenverarbeitung - Übersicht", help="Zeigt alle geladenen Dateien und deren Status. Prüfen Sie, ob alle benötigten Dateien hochgeladen wurden, bevor Sie die Verarbeitung starten.")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.write("**Leitdatei:**")
            for anmeldeart in VERARBEITBARE_ARTEN + PAUSCHALE_ARTEN:
                if st.session_state.stats.get(anmeldeart, 0) > 0:
                    st.write(f"{anmeldeart}: {st.session_state.stats.get(anmeldeart, 0)}")
            if st.session_state.stats.get('(leer)', 0) > 0:
                st.write(f"(leer): {st.session_state.stats.get('(leer)', 0)}")
        
        with col2:
            st.write("**Import EZA:**")
            if st.session_state.df_import_eza is not None:
                st.write(f"{len(st.session_state.df_import_eza)} Einträge")
            else:
                st.write("—")
        
        with col3:
            st.write("**Import ZL:**")
            if st.session_state.df_import_zl is not None:
                st.write(f"{len(st.session_state.df_import_zl)} Einträge")
            else:
                st.write("—")
        
        with col4:
            st.write("**NCTS:**")
            if is_dataframe_valid(st.session_state.get('df_ncts')):
                st.write(f"{len(st.session_state.df_ncts)} Einträge")
            else:
                st.write("—")
        
        if st.session_state.get('ncar_enabled', True):
            st.write("")
            if 'df_ncar' in st.session_state and st.session_state['df_ncar'] is not None:
                st.info(f"✨ NCAR-Datei: {len(st.session_state['df_ncar'])} Einträge")
            else:
                st.info("✨ NCAR-Datei: Nicht hochgeladen (optional)")
        
        if st.session_state.get('buergschaft_erhöhung_aktiv', False):
            betrag = st.session_state.get('buergschaft_erhöhung_betrag', 0)
            datum = st.session_state.get('buergschaft_erhöhung_datum', date.today())
            st.info(f"💰 Bürgschaftserhöhung: +{betrag:,.0f} € am {datum.strftime('%d.%m.%Y')} aktiviert")
        
        can_process = True
        missing_files = []
        
        if st.session_state.stats.get("IMDC", 0) > 0 and st.session_state.df_import_eza is None:
            can_process = False
            missing_files.append("Import EZA")
        
        if st.session_state.stats.get("WIDS", 0) > 0 and st.session_state.df_import_zl is None:
            can_process = False
            missing_files.append("Import ZL")
        
        if st.session_state.stats.get("NCDP", 0) > 0 and not is_dataframe_valid(st.session_state.get('df_ncts')):
            can_process = False
            missing_files.append("NCTS")
        
        if not can_process:
            st.error(f"Fehlende Dateien: {', '.join(missing_files)}")
        else:
            st.markdown("")
            if st.button("🚀 Verarbeitung starten", use_container_width=True,
                        help="Startet die Berechnung der Bürgschaftsbelastung mit allen hochgeladenen Dateien"):
                st.info("🔄 Verarbeitung wurde gestartet...")
                process_data()

# === HISTORIE FUNKTIONEN ===

def get_history_file():
    """Gibt den Pfad zur Historie-Datei zurück"""
    if 'mandant' not in st.session_state:
        return None
    
    mandant_clean = st.session_state['mandant'].lower().replace(' ', '_')
    return f"historie_{mandant_clean}.json"

def load_history():
    """Lädt die Historie aus der Datei"""
    history_file = get_history_file()
    if not history_file or not os.path.exists(history_file):
        return []
    
    try:
        with open(history_file, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return []

def save_history(history):
    """Speichert die Historie in die Datei"""
    history_file = get_history_file()
    if not history_file:
        return
    
    try:
        with open(history_file, 'w', encoding='utf-8') as f:
            json.dump(history, f, indent=2, ensure_ascii=False)
    except Exception as e:
        st.error(f"❌ Fehler beim Speichern der Historie: {e}")

def save_to_history():
    """Speichert die aktuelle Verarbeitung in der Historie"""
    if not all([
        'excel_file' in st.session_state,
        'von_datum' in st.session_state,
        'bis_datum' in st.session_state,
        'ziel_sorted' in st.session_state
    ]):
        return
    
    import base64
    
    history = load_history()
    
    # Erstelle Historie-Eintrag
    entry = {
        'timestamp': datetime.now().isoformat(),
        'von_datum': st.session_state['von_datum'].strftime('%d.%m.%Y'),
        'bis_datum': st.session_state['bis_datum'].strftime('%d.%m.%Y'),
        'zeilen': len(st.session_state['ziel_sorted']),
        'startbuergschaft': st.session_state.get('startbuergschaft', 0),
        'max_auslastung': st.session_state.get('max_auslastung', 'N/A'),
        'excel_data': base64.b64encode(st.session_state['excel_file']).decode('utf-8'),
        'stats': st.session_state.get('stats', {}),
        'processing_stats': st.session_state.get('processing_stats', {}),
        'config': {
            'pauschalbetrag': st.session_state.get('pauschalbetrag', 10000),
            'ersatz_zollsatz': st.session_state.get('zollsatz_ersatz', 0.12),
            'buergschaft_erhöhung_aktiv': st.session_state.get('buergschaft_erhöhung_aktiv', False),
            'buergschaft_erhöhung_datum': st.session_state.get('buergschaft_erhöhung_datum', date.today()).strftime('%d.%m.%Y') if st.session_state.get('buergschaft_erhöhung_aktiv', False) else None,
            'buergschaft_erhöhung_betrag': st.session_state.get('buergschaft_erhöhung_betrag', 0) if st.session_state.get('buergschaft_erhöhung_aktiv', False) else None
        }
    }
    
    # Füge am Anfang hinzu
    history.insert(0, entry)
    
    # Behalte nur die letzten 30 Einträge
    history = history[:30]
    
    save_history(history)

def show_history_page():
    """Zeigt die Historie-Seite mit verbesserter Radio-Button-Tabellen-Ansicht"""
    st.title("📊 Bürgschaftsverlauf")
    
    history = load_history()
    
    if not history:
        st.info("ℹ️ Noch keine Historie vorhanden.")
        st.write("Nach der ersten erfolgreichen Verarbeitung werden hier alle Ihre Verarbeitungen gespeichert.")
        return
    
    st.success(f"✅ {len(history)} gespeicherte Verarbeitungen")
    
    # Erstelle 40-60 Layout
    col_left, col_right = st.columns([2, 3])
    
    with col_left:
        st.subheader("📋 Verarbeitungen auswählen")
        st.caption("Wählen Sie einen Eintrag aus der Liste")
        
        # Erstelle Header
        st.markdown("""
        <div style="font-family: monospace; font-size: 0.9em; color: #666; margin-bottom: 5px;">
        Nr. | Datum      | Zeit  | Bürgschaftszeitraum
        -------------------------------------------------
        </div>
        """, unsafe_allow_html=True)
        
        # Erstelle formatierte Optionen für Radio-Buttons
        options = []
        for i, entry in enumerate(history):
            timestamp = datetime.fromisoformat(entry['timestamp'])
            
            # Formatiere mit festen Breiten
            nr = f"{i+1:02d}"  # Immer 2-stellig
            datum = timestamp.strftime('%d.%m.%Y')  # Immer 10 Zeichen
            zeit = timestamp.strftime('%H:%M')  # Immer 5 Zeichen
            von_dt = datetime.strptime(entry['von_datum'], '%d.%m.%Y')
            bis_dt = datetime.strptime(entry['bis_datum'], '%d.%m.%Y')
            zeitraum = f"{von_dt.strftime('%m-%Y')} bis {bis_dt.strftime('%m-%Y')}"
            
            # Verwende Pipes als Trenner
            option = f"{nr} | {datum} | {zeit} | {zeitraum}"
            options.append(option)
        
        # Radio-Buttons mit formatierter Liste
        selected = st.radio(
            "Auswahl",
            options,
            label_visibility="collapsed",
            key="history_selection"
        )
        
        # Zusatz-Info unter der Auswahl
        if selected:
            index = options.index(selected)
            entry = history[index]
            st.caption(f"📊 {entry['zeilen']} Zeilen | 💰 {format_currency(entry['startbuergschaft'])} | ⚡ {entry['max_auslastung']}")
    
    with col_right:
        if selected:
            index = options.index(selected)
            entry = history[index]
            
            st.subheader("📄 Details zur ausgewählten Verarbeitung")
            
            # Basis-Informationen
            col1, col2 = st.columns(2)
            with col1:
                st.metric("Zeitraum", f"{entry['von_datum']} - {entry['bis_datum']}")
                st.metric("Verarbeitete Zeilen", format_currency(entry['zeilen'], display_only=True))
            with col2:
                st.metric("Bürgschaft", format_currency(entry['startbuergschaft']))
                st.metric("Max. Auslastung", entry['max_auslastung'])
            
            st.markdown("---")
            
            # Download-Buttons
            st.subheader("📥 Downloads")
            
            col1, col2 = st.columns(2)
            
            import base64
            
            with col1:
                excel_data = base64.b64decode(entry['excel_data'])
                
                mandant_prefix = st.session_state.get('mandant', 'Unbekannt')[:3].upper()
                von_dt = datetime.strptime(entry['von_datum'], '%d.%m.%Y')
                bis_dt = datetime.strptime(entry['bis_datum'], '%d.%m.%Y')
                excel_filename = f"Verwahrliste_{mandant_prefix}_{von_dt.strftime('%m_%y')}#{bis_dt.strftime('%m_%y')}_Historie.xlsx"
                
                st.download_button(
                    label="📊 Excel herunterladen",
                    data=excel_data,
                    file_name=excel_filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    type="primary"
                )
            
            with col2:
                if st.button("📄 Dokumentation erstellen", use_container_width=True, type="secondary"):
                    # Lade die alten Werte in den Session State für die Doku-Erstellung
                    st.session_state['von_datum'] = datetime.strptime(entry['von_datum'], '%d.%m.%Y').date()
                    st.session_state['bis_datum'] = datetime.strptime(entry['bis_datum'], '%d.%m.%Y').date()
                    st.session_state['startbuergschaft'] = entry['startbuergschaft']
                    st.session_state['stats'] = entry['stats']
                    st.session_state['processing_stats'] = entry.get('processing_stats', {})
                    st.session_state['max_auslastung'] = entry['max_auslastung']
                    st.session_state['max_auslastung_str'] = entry['max_auslastung'].replace('.', ',')
                    st.session_state['pauschalbetrag'] = entry['config']['pauschalbetrag']
                    st.session_state['zollsatz_ersatz'] = entry['config']['ersatz_zollsatz']
                    
                    if 'ziel_sorted' not in st.session_state:
                        st.session_state['ziel_sorted'] = pd.DataFrame()
                    
                    # Generiere Doku
                    with st.spinner("Dokumentation wird erstellt..."):
                        doc_data = create_personalized_documentation()
                        if doc_data:
                            st.session_state['doc_data_history'] = doc_data
                            st.session_state['doc_filename_history'] = f"Zoll_Dokumentation_{mandant_prefix}_{von_dt.strftime('%m_%y')}#{bis_dt.strftime('%m_%y')}_Historie.docx"
                            st.rerun()
                        else:
                            st.error("❌ Dokumentation konnte nicht erstellt werden")
            
            # Zeige Download-Button für erstellte Dokumentation
            if 'doc_data_history' in st.session_state and st.session_state.get('doc_data_history'):
                st.download_button(
                    label="📥 Word-Dokument herunterladen",
                    data=st.session_state['doc_data_history'],
                    file_name=st.session_state.get('doc_filename_history', 'Zoll_Dokumentation_Historie.docx'),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                # Lösche nach Download
                del st.session_state['doc_data_history']
                del st.session_state['doc_filename_history']
            
            # Erweiterte Details im Expander
            with st.expander("📊 Erweiterte Details anzeigen", expanded=False):
                # Konfiguration
                if 'config' in entry:
                    st.subheader("⚙️ Verwendete Konfiguration")
                    config = entry['config']
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        st.write(f"**Pauschalbetrag:** {format_currency(config['pauschalbetrag'])}")
                        st.write(f"**Ersatz-Zollsatz:** {config['ersatz_zollsatz'] * 100:.0f}%")
                    
                    with col2:
                        if config.get('buergschaft_erhöhung_aktiv', False):
                            st.write(f"**Bürgschaftserhöhung:** {format_currency(config['buergschaft_erhöhung_betrag'])}")
                            st.write(f"**Erhöhung am:** {config['buergschaft_erhöhung_datum']}")
                        else:
                            st.write("**Bürgschaftserhöhung:** Keine")
                
                # Verarbeitungsstatistiken
                if 'stats' in entry:
                    st.markdown("---")
                    st.subheader("📈 Verarbeitungsstatistiken")
                    
                    stats = entry['stats']
                    verarbeitbare = ['IMDC', 'WIDS', 'IPDC', 'NCDP', '(leer)', 'APDC', 'AVDC', 'NCAR']
                    
                    stats_data = []
                    for art in verarbeitbare:
                        if stats.get(art, 0) > 0:
                            stats_data.append({
                                'Anmeldeart': art,
                                'Anzahl': stats.get(art, 0)
                            })
                    
                    if stats_data:
                        df_stats = pd.DataFrame(stats_data)
                        st.dataframe(df_stats, hide_index=True, use_container_width=True)
                
                # Processing Stats
                if 'processing_stats' in entry and entry['processing_stats']:
                    st.markdown("---")
                    st.subheader("🔧 Verarbeitungsprotokoll")
                    
                    proc_stats = entry['processing_stats']
                    protocol_data = []
                    
                    if proc_stats.get('processed_imdc', 0) > 0:
                        protocol_data.append({
                            'Anmeldeart': 'IMDC',
                            'Verarbeitet': proc_stats.get('processed_imdc', 0),
                            'Mit Match': proc_stats.get('imdc_match', 0),
                            'Ohne Match': proc_stats.get('imdc_no_match', 0)
                        })
                    
                    if proc_stats.get('processed_wids', 0) > 0:
                        protocol_data.append({
                            'Anmeldeart': 'WIDS',
                            'Verarbeitet': proc_stats.get('processed_wids', 0),
                            'Mit Match': proc_stats.get('wids_match', 0),
                            'Ohne Match': proc_stats.get('wids_no_match', 0)
                        })
                    
                    if proc_stats.get('processed_ncdp', 0) > 0:
                        protocol_data.append({
                            'Anmeldeart': 'NCDP',
                            'Verarbeitet': proc_stats.get('processed_ncdp', 0),
                            'Mit Match': proc_stats.get('ncdp_match', 0),
                            'Ohne Match': proc_stats.get('ncdp_no_match', 0)
                        })
                    
                    if protocol_data:
                        df_protocol = pd.DataFrame(protocol_data)
                        st.dataframe(df_protocol, hide_index=True, use_container_width=True)
        else:
            st.info("👈 Wählen Sie links einen Eintrag aus, um Details anzuzeigen")

# === HAUPTFUNKTION ===

def main():
    """Hauptfunktion der App"""
    show_login()
    
    if st.session_state.get('authenticated', False):
        if not check_initial_setup():
            return
        
        init_session_state()
        setup_sidebar()
        
        has_downloads = 'excel_file' in st.session_state and st.session_state['excel_file'] is not None
        
        # HISTORIE-HOOK: Historie-Tab hinzufügen
        selected = option_menu(
            menu_title=None,
            options=["Verarbeitung", "Downloads", "Bürgschaftsverlauf", "Einstellungen"],
            icons=["clipboard-data", "download", "clock-history", "gear"],
            default_index=1 if has_downloads and st.session_state.get('show_downloads', False) else 0,
            orientation="horizontal",
            styles={
                "container": {"padding": "0!important", "background-color": "#fafafa"},
                "icon": {"color": "#8B1C1C", "font-size": "18px"},
                "nav-link": {
                    "font-size": "16px", 
                    "text-align": "center", 
                    "margin": "0px",
                    "--hover-color": "#fee2e2"
                },
                "nav-link-selected": {
                    "background-color": "#8B1C1C",
                    "color": "white"
                },
            }
        )
        
        if has_downloads and selected == "Downloads":
            st.success("✅ Downloads bereit!")
        
        if selected == "Verarbeitung":
            show_file_upload_section()
        elif selected == "Downloads":
            st.session_state.show_downloads = True
            show_downloads_section()
        elif selected == "Bürgschaftsverlauf":
            show_history_page()
        elif selected == "Einstellungen":
            show_settings_page()

if __name__ == "__main__":
    main()